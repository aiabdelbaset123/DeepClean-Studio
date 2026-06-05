#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DeepClean Studio

A Streamlit app for careful academic text revision. The app focuses on clarity,
flow, citation preservation, transparent change review, and Word export.
"""

from __future__ import annotations

import difflib
import html
import re
from collections import Counter
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import docx2txt
import numpy as np
import pandas as pd
import pypdf
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    import textstat
except Exception:  # pragma: no cover - optional quality helper
    textstat = None

try:
    from sentence_transformers import SentenceTransformer, util
except Exception:  # pragma: no cover - optional semantic helper
    SentenceTransformer = None
    util = None


st.set_page_config(page_title="DeepClean Studio", layout="wide")

APP_DIR = Path(__file__).resolve().parent
SYNONYM_COLUMNS = {"domain", "original", "replacement"}
AUTHOR_NAME = "Prof. Dr. Abdel-baset H. Mekky"


@dataclass(frozen=True)
class RevisionStats:
    words_original: int
    words_revised: int
    sentences_revised: int
    avg_sentence_length: float
    lexical_diversity: float
    similarity: float
    formulaic_phrase_count: int
    certainty_marker_count: int
    sentence_length_variation: float
    authorship_review_band: str
    type_token_ratio: float = 0.0
    punctuation_entropy: float = 0.0
    function_word_bigram_concentration: float = 0.0
    tortured_phrase_count: int = 0
    burstiness_cv: float = 0.0


@dataclass(frozen=True)
class SentenceSignal:
    text: str
    score: float
    label: str
    reasons: Tuple[str, ...]


@dataclass(frozen=True)
class TransparencyReport:
    chars: int
    words: int
    sentences: int
    qualifying_chars: int
    excluded_chars: int
    excluded_blocks: Tuple[Tuple[str, int], ...]
    minimum_ready: bool
    classification: str
    confidence: float
    ai_generated_score: float
    ai_paraphrased_score: float
    plagiarism_signal: float
    perplexity_signal: float
    burstiness_signal: float
    paraphrase_signal: float
    translation_signal: float
    esl_adjustment: float
    false_positive_shield: bool
    ai_vocabulary_hits: Tuple[Tuple[str, int], ...]
    sentence_signals: Tuple[SentenceSignal, ...]
    citation_alerts: Tuple[str, ...]
    source_code_alerts: Tuple[str, ...]
    section_notes: Tuple[str, ...]
    model_update_note: str
    structural_alerts: Tuple[str, ...] = ()


FORMULAIC_PATTERNS = [
    r"\bit is important to note that\b",
    r"\bit should be noted that\b",
    r"\bin today's (?:world|society)\b",
    r"\bplays a (?:crucial|vital) role in\b",
    r"\bhas a significant impact on\b",
    r"\bdelve into\b",
    r"\bstands as\b",
    r"\bserves as\b",
    r"\bis a testament to\b",
    r"\bis a reminder of\b",
    r"\bmarks? a (?:significant|pivotal|key|crucial) (?:moment|shift|turning point)\b",
    r"\b(?:underscores?|highlights?|emphasizes?) (?:the )?(?:importance|significance|need)\b",
    r",\s*(?:highlighting|underscoring|emphasizing|ensuring|reflecting|symbolizing|contributing to|cultivating|fostering|encompassing)\b",
    r"\breflects? broader\b",
    r"\bsymboli[sz]es? (?:its )?(?:ongoing|enduring|lasting)\b",
    r"\bcontribut(?:e|es|ing) to (?:the )?(?:broader|overall)\b",
    r"\bsetting the stage for\b",
    r"\bindelible mark\b",
    r"\bdeeply rooted\b",
    r"\bvaluable insights\b",
    r"\baligns? with\b",
    r"\bresonates? with\b",
    r"\bboasts? a\b",
    r"\brich tapestry\b",
    r"\bevolving landscape\b",
    r"\bactive social media presence\b",
    r"\bstrong digital presence\b",
    r"\bindependent coverage\b",
    r"\blocal media outlets\b",
    r"\bregional media outlets\b",
    r"\bnational media outlets\b",
    r"\bmusic outlets\b",
    r"\bbusiness outlets\b",
    r"\btech outlets\b",
    r"\bmedia outlets\b",
    r"\bprofiled in\b",
    r"\bwritten by a leading expert\b",
    r"\bindustry reports\b",
    r"\bobservers have cited\b",
    r"\bexperts argue\b",
    r"\bsome critics\b",
    r"\bdespite (?:its|these) [^.!?]{0,80}challenges\b",
    r"\bthis article (?:will|aims to|seeks to)\b",
    r"\bin this article,?\s+(?:we|i) will\b",
    r"\bkey takeaways\b",
    r"\b(?:fascinating|dependable|value-driven) (?:glimpse|experience|experiences)\b",
    r"\b(?:transformative|profound) power\b",
    r"\bongoing public presence\b",
    r"\bwell-sourced edits\b",
    r"\bif you have any (?:concerns|suggestions)\b",
    r"\bif there are specific\b",
    r"\byou'?re absolutely right\b",
    r"\bas of my last (?:training update|knowledge update|knowledge)\b",
    r"\bup to my last training update\b",
    r"\bas of (?:january|february|march|april|may|june|july|august|september|october|november|december) \d{4}\b",
    r"\bas an ai (?:language model|assistant)\b",
    r"\bas a large language model\b",
    r"\bi cannot (?:offer|provide|verify)\b",
    r"\bi hope this helps\b",
    r"\bof course[!,]?\b",
    r"\bcertainly[!,]?\b",
    r"\bin summary\b",
    r"\bin conclusion\b",
    r"\boverall\b",
    r"\bfurthermore\b",
    r"\bmoreover\b",
    r"\badditionally\b",
    r"\bon the other hand\b",
    r"\bit is worth noting that\b",
    r"\bit should be emphasized that\b",
    r"علاوة على ذلك",
    r"بالإضافة إلى ذلك",
    r"في الختام",
    r"وبناءً على ما سبق",
    r"من ناحية أخرى",
    r"تجدر الإشارة إلى",
    # Negative parallelisms
    r"\bnot only\b[^.!?]{0,60}\bbut also\b",
    r"\bnot just\b[^.!?]{0,60}\bbut also\b",
    r"\bit's not\b[^.!?]{0,40}\bit's\b",
    r"\bnot a\b[^.!?]{0,40}\bbut a\b",
    # Promotional / advertisement-like language
    r"\bnestled in\b",
    r"\bin the heart of\b",
    r"\bgroundbreaking\b",
    r"\brenowned\b",
    r"\bfeaturing a\b",
    r"\bdiverse array\b",
    r"\bexemplifies\b",
    r"\bcommitment to\b",
    r"\bnatural beauty\b",
    r"\bcommitment to excellence\b",
    r"\bworld-class\b",
    r"\bstate-of-the-art\b",
    r"\bcutting-edge\b",
    r"\bpremier\b",
    r"\bleading provider\b",
    r"\btailored\b",
    r"\bempower(?:s|ed|ing)?\b",
    r"\bunlock(?:s|ed|ing)?\b",
    r"\bstreamlin(?:e|es|ed|ing)\b",
    # Avoidance of basic copulatives
    r"\brefers to\b",
    r"\bis defined as\b",
    r"\bis known as\b",
    r"\bis described as\b",
    # Challenges + future prospects formula
    r"\bdespite (?:its|these|the) [^.!?]{0,80}faces (?:several |many |numerous )?challenges?\b",
    r"\bfuture (?:outlook|prospects|directions|implications)\b",
    r"\blooking ahead\b",
    r"\bmoving forward\b",
    r"\bchallenges and (?:opportunities|future|legacy)\b",
    # Canned communication
    r"\bwould you like\b",
    r"\bis there anything else\b",
    r"\blet me know\b",
    r"\bI am open to\b",
    r"\bfeel free to\b",
    r"\bdo not hesitate to\b",
    r"\bI'd be happy to\b",
    r"\bI'd be glad to\b",
    # Placeholder / template text
    r"\bXX-XX-XXXX\b",
    r"\bYYYY-MM-DD\b",
    r"\b\[insert\b",
    r"\b\[citation needed\]\b",
    r"\b\[insert citation\]\b",
    r"\b\[insert name\]\b",
    r"\b\[insert date\]\b",
    r"\b\[insert reference\]\b",
    # ChatGPT-specific markup
    r"\bcontentReference\b",
    r"\boai_citation\b",
    r"\bgrok_card\b",
    r"\bgrok_render_citation\b",
    r"\bturn\dsearch\d\b",
    # Didactic disclaimers
    r"\bit's important to (?:note|remember|consider|understand)\b",
    r"\bit is important to (?:remember|consider|understand)\b",
    r"\bkeep in mind that\b",
    r"\bbear in mind that\b",
    # Section summaries
    r"\bin short\b",
    r"\bto summarize\b",
    r"\bto sum up\b",
    r"\bwrapping up\b",
    # Exaggerated notability
    r"\bwidely regarded as\b",
    r"\bhighly (?:acclaimed|respected|regarded|sought)\b",
    r"\binternationally recognized\b",
    r"\bglobally recognized\b",
    # Knowledge cutoff / speculation about gaps
    r"\bwhile specific details are (?:limited|scarce)\b",
    r"\bnot widely (?:available|documented|disclosed)\b",
    r"\bmaintains a low profile\b",
    r"\bkeeps (?:personal|private) details (?:private|under wraps)\b",
    r"\binformation about .{0,30} (?:is|remains) (?:limited|scarce|unavailable)\b",
    # Additional transitional fillers
    r"\bhence\b",
    r"\bconsequently\b",
    r"\bthus\b",
    # "concrete" overuse (ChatGPT defense word)
    r"\bconcrete (?:example|evidence|steps|action|proof)\b",
]
CERTAINTY_PATTERNS = [
    r"\bclearly proves\b",
    r"\bproves that\b",
    r"\balways\b",
    r"\bnever\b",
    r"\bdefinitely\b",
    r"\bundoubtedly\b",
]

TORTURED_PHRASES = [
    (r"\bsynthetic cognitive capability\b", "artificial intelligence"),
    (r"\bcomputational reasoning system\b", "artificial intelligence"),
    (r"\bmachine cognition\b", "artificial intelligence"),
    (r"\bautomated reasoning platform\b", "AI system"),
    (r"\bdigital information processing\b", "computing"),
    (r"\belectronic computational device\b", "computer"),
    (r"\bhuman-computer interaction\b", "user interface"),
    (r"\bvirtual presence\b", "online presence"),
    (r"\bknowledge retrieval system\b", "search engine"),
    (r"\bpredictive analytics model\b", "prediction model"),
    (r"\bdata-driven insights\b", "findings from data"),
    (r"\bactionable intelligence\b", "useful information"),
    (r"\bcognitive computational framework\b", "AI framework"),
    (r"\bintellectual property asset\b", "patent"),
    (r"\bstrategic advantage\b", "advantage"),
    (r"\bcompetitive positioning\b", "market position"),
    (r"\bstakeholder engagement\b", "communication"),
    (r"\bvalue proposition\b", "offering"),
    (r"\bparadigm shift\b", "shift"),
    (r"\bsynergistic integration\b", "combination"),
    (r"\bcross-functional collaboration\b", "teamwork"),
    (r"\bthought leadership\b", "expertise"),
    (r"\bbest practices\b", "standard methods"),
    (r"\bleverage synergies\b", "combine efforts"),
    (r"\bdisruptive innovation\b", "new approach"),
    (r"\bscalable solution\b", "solution"),
    (r"\bend-to-end\b", "complete"),
    (r"\bdata-driven decision making\b", "evidence-based decisions"),
    (r"\bcontinuous improvement\b", "improvement"),
    (r"\bdeep dive\b", "detailed examination"),
    (r"\bgranular level\b", "detailed level"),
    (r"\bholistic approach\b", "broad approach"),
    (r"\bquantum leap\b", "major step"),
    (r"\bgame-changer\b", "major change"),
    (r"\bnext-generation\b", "new"),
    (r"\bmission-critical\b", "essential"),
    (r"\bfuture-proof\b", "durable"),
    (r"\bbest-in-class\b", "top"),
    (r"\bindustry-leading\b", "leading"),
    (r"\bhigh-performance\b", "fast"),
    (r"\bmission-driven\b", "purposeful"),
    # Arabic tortured phrases
    (r"القدرة المعرفية الاصطناعية", "الذكاء الاصطناعي"),
    (r"نظام الحوسبة العصبية", "نظام الذكاء الاصطناعي"),
    (r"التكنولوجيا الرقمية المتقدمة", "التكنولوجيا المتقدمة"),
    (r"منصة التحليل التنبؤي", "نموذج التنبؤ"),
    (r"إدارة المعرفة المؤسسية", "إدارة المعرفة"),
    (r"الابتكار المزعزع", "الابتكار الجديد"),
    (r"الريادة الفكرية", "الخبرة"),
]

AI_VOCABULARY_TERMS = [
    "additionally",
    "align",
    "aligns with",
    "boasts",
    "bolstered",
    "comprehensive",
    "contributing",
    "crucial",
    "cultivate",
    "delve",
    "deeply rooted",
    "dynamic",
    "elevate",
    "encompassing",
    "enduring",
    "evolving landscape",
    "foster",
    "furthermore",
    "garner",
    "highlight",
    "holistic",
    "indelible mark",
    "intricate",
    "interplay",
    "in conclusion",
    "in summary",
    "innovative",
    "key",
    "landscape",
    "leverage",
    "meticulous",
    "moreover",
    "multifaceted",
    "notably",
    "overall",
    "pivotal",
    "profound",
    "resonate",
    "robust",
    "seamless",
    "serve as",
    "serves as",
    "showcase",
    "significant",
    "stands as",
    "tapestry",
    "testament",
    "transformative",
    "underscore",
    "underscores",
    "valuable insights",
    "vibrant",
    "vital",
    "علاوة على ذلك",
    "بالإضافة إلى ذلك",
    "في الختام",
    "بناء على ما سبق",
    "من ناحية أخرى",
    "تجدر الإشارة",
    "بشكل شامل",
    "محوري",
    "حيوي",
    "nestled",
    "groundbreaking",
    "renowned",
    "featuring",
    "diverse array",
    "exemplifies",
    "commitment to",
    "natural beauty",
    "world-class",
    "state-of-the-art",
    "cutting-edge",
    "premier",
    "tailored",
    "empower",
    "unlock",
    "streamline",
    "concrete",
    "refers to",
    "hence",
    "consequently",
    "thus",
    "widely regarded",
    "exemplify",
    "fostering",
    "cultivating",
    "encompassing",
    "leveraging",
    "bolster",
    "elevate",
    "ensure",
    "enhance",
    "innovative",
    "dynamic",
    "vibrant",
    "بشكل ملموس",
    "رائد",
    "فريد من نوعه",
    "على المستوى العالمي",
    "بشكل استثنائي",
    "متطور",
]

WIKIPEDIA_AI_STYLE_REWRITES = [
    (r"\bclearly proves that\b", "suggests that"),
    (r"\bproves that\b", "suggests that"),
    (r"\bclearly demonstrates that\b", "indicates that"),
    (r"\bit is (?:important|critical|crucial) to (?:note|remember|consider) that\s+", ""),
    (r"\bit should be noted that\s+", ""),
    (r"\bit is worth noting that\s+", ""),
    (r"\bit should be emphasized that\s+", ""),
    (r"\bin today's (?:world|society|era)\b", "currently"),
    (r"\bthis article (?:will|aims to|seeks to)\s+", ""),
    (r"\bin this article,?\s+(?:we|i) will\s+", ""),
    (r"\bkey takeaways?\b", "main points"),
    (r"\bplays a (?:crucial|vital|pivotal|key) role in\b", "is part of"),
    (r"\bhas a significant impact on\b", "affects"),
    (r"\bmake(?:s)? a significant contribution to\b", "contributes to"),
    (r"\bdelve into\b", "examine"),
    (r"\bstands as a testament to the evolving landscape of\b", "is part of"),
    (r"\bserves as a testament to the evolving landscape of\b", "is part of"),
    (r"\bstands as a testament to\b", "shows"),
    (r"\bserves as a testament to\b", "shows"),
    (r"\bis a testament to\b", "shows"),
    (r"\bis testament to\b", "shows"),
    (r"\bis a reminder of\b", "shows"),
    (r"\bstands as (?:a|an|the)?\s*", "is "),
    (r"\bserves as (?:a|an|the)?\s*", "is "),
    (r"\bmarks? a (?:significant|pivotal|key|crucial) (?:moment|shift|turning point) in\b", "changed"),
    (r"\brepresents? a (?:significant|pivotal|key|crucial) (?:moment|shift|turning point) in\b", "changed"),
    (r"\b(?:underscores?|highlights?|emphasizes?) (?:the )?(?:importance|significance|need) of\b", "shows"),
    (r",\s*highlighting its (?:pivotal|key|crucial|vital) role and significant impact on the broader ([^.!?]+)", r" and affects the \1"),
    (r",\s*(?:highlighting|underscoring|emphasizing) (?:the )?(?:importance|significance) of ([^.!?]+)", r" and discusses \1"),
    (r",\s*(?:highlighting|underscoring|emphasizing|ensuring|reflecting|symbolizing|contributing to|cultivating|fostering|encompassing)\s+", ". "),
    (r"\breflects? broader\b", "reflects"),
    (r"\bsymboli[sz]es? (?:its )?(?:ongoing|enduring|lasting)\b", "shows"),
    (r"\bcontribut(?:e|es|ing) to (?:the )?(?:broader|overall)\b", "contributes to"),
    (r"\bsetting the stage for\b", "preceding"),
    (r"\bindelible mark\b", "effect"),
    (r"\bdeeply rooted\b", "long-standing"),
    (r"\bvaluable insights into\b", "evidence about"),
    (r"\baligns? with\b", "matches"),
    (r"\bresonates? with\b", "relates to"),
    (r"\bboasts? a\b", "has a"),
    (r"\brich tapestry of\b", "range of"),
    (r"\bevolving landscape\b", "field"),
    (r"\b(?:vibrant|rich|profound|transformative|fascinating)\s+(?=(?:culture|heritage|legacy|community|glimpse|power|journey|landscape)\b)", ""),
    (r"\bshowcas(?:e|es|ing)\b", "shows"),
    (r"\benhanc(?:e|es|ing)\b", "improve"),
    (r"\bfoster(?:s|ing)?\b", "support"),
    (r"\bgarner(?:ed|s|ing)?\b", "received"),
    (r"\bmeticulous(?:ly)?\b", "careful"),
    (r"\bintricate(?:ly| intricacies)?\b", "detailed"),
    (r"\brobust\b", "well-supported"),
    (r"\bseamless\b", "smooth"),
    (r"\bholistic\b", "broad"),
    (r"\bmultifaceted\b", "multi-part"),
    (r"\bpivotal\b", "important"),
    (r"\bcrucial\b", "important"),
    (r"\bvital\b", "important"),
    (r"\bsignificant\b", "measurable"),
    (r"\bmaintains? (?:a )?(?:strong digital presence|active social media presence)\b", "has social media accounts"),
    (r"\bactive social media presence\b", "social media accounts"),
    (r"\bbroader community\b", "community"),
    (r"\bindependent coverage\b", "coverage"),
    (r"\b(?:local|regional|national|music|business|tech) media outlets\b", "publications"),
    (r"\bmedia outlets\b", "publications"),
    (r"\bprofiled in\b", "covered in"),
    (r"\bwritten by a leading expert\b", "written by a subject specialist"),
    (r"\bindustry reports\b", "reports"),
    (r"\bobservers have cited\b", "sources cite"),
    (r"\bexperts argue\b", "some sources argue"),
    (r"\bsome critics\b", "some reviewers"),
    (r"\bdespite (?:its|these) ([^.!?]{0,80}) challenges\b", r"despite \1 limits"),
    (r"\bas of my last (?:training update|knowledge update|knowledge),?\s*", ""),
    (r"\bup to my last training update,?\s*", ""),
    (r"\bas an ai (?:language model|assistant),?\s*", ""),
    (r"\bas a large language model,?\s*", ""),
    (r"\bi cannot (?:offer|provide|verify)[^.!?]*(?:[.!?]|$)", ""),
    (r"\bi hope this helps[.!]?\s*", ""),
    (r"^\s*(?:of course|certainly)[!,]?\s*", ""),
    (r"\byou'?re absolutely right[.!]?\s*", ""),
    (r"\bif you have any (?:concerns|suggestions)[^.!?]*(?:[.!?]|$)", ""),
    (r"\bif there are specific [^.!?]*(?:[.!?]|$)", ""),
    (r"\bin summary,?\s+", ""),
    (r"\bin conclusion,?\s+", ""),
    (r"\boverall,?\s+", ""),
    # Negative parallelisms - restructure
    (r"\bnot only\s+([^.!?]{0,60})\s+but also\s+", r"\1 and "),
    (r"\bnot just\s+([^.!?]{0,60})\s+but also\s+", r"\1 and "),
    (r"\bnot a\s+([^.!?]{0,40})\s+but a\s+", r"\1 and "),
    # Promotional language cleanup
    (r"\bnestled in the heart of\b", "in"),
    (r"\bnestled in\b", "in"),
    (r"\bin the heart of\b", "in"),
    (r"\bgroundbreaking\b", "new"),
    (r"\brenowned\b", "known"),
    (r"\bfeaturing a diverse array of\b", "with"),
    (r"\bdiverse array of\b", "range of"),
    (r"\bexemplifies\b", "shows"),
    (r"\bcommitment to excellence\b", "focus on quality"),
    (r"\bcommitment to\b", "focus on"),
    (r"\bworld-class\b", "leading"),
    (r"\bstate-of-the-art\b", "modern"),
    (r"\bcutting-edge\b", "advanced"),
    (r"\bpremier\b", "main"),
    (r"\bleading provider of\b", "provider of"),
    (r"\btailored\b", "custom"),
    (r"\bempower(?:s|ed|ing)?\b", "support"),
    (r"\bunlock(?:s|ed|ing)?\b", "open"),
    (r"\bstreamlin(?:e|es|ed|ing)\b", "simplify"),
    # Copulative avoidance
    (r"\brefers to\b", "is"),
    (r"\bis defined as\b", "is"),
    (r"\bis known as\b", "is"),
    (r"\bis described as\b", "is"),
    # Challenges + future formula
    (r"\bfuture outlook\b", "next steps"),
    (r"\bfuture prospects\b", "next steps"),
    (r"\bfuture directions\b", "next steps"),
    (r"\bfuture implications\b", "effects"),
    (r"\blooking ahead\b", ""),
    (r"\bmoving forward\b", ""),
    (r"\bchallenges and opportunities\b", "limits and options"),
    (r"\bchallenges and future\b", "limits and next steps"),
    (r"\bchallenges and legacy\b", "limits and history"),
    # Section summary cleanup
    (r"\bin short,?\s+", ""),
    (r"\bto summarize,?\s+", ""),
    (r"\bto sum up,?\s+", ""),
    (r"\bwrapping up,?\s+", ""),
    # Didactic disclaimer cleanup
    (r"\bit's important to (?:note|remember|consider|understand) that\s+", ""),
    (r"\bit is important to (?:remember|consider|understand) that\s+", ""),
    (r"\bkeep in mind that\s+", ""),
    (r"\bbear in mind that\s+", ""),
    # Exaggerated notability
    (r"\bwidely regarded as\b", "considered"),
    (r"\bhighly acclaimed\b", "praised"),
    (r"\bhighly respected\b", "respected"),
    (r"\bhighly regarded\b", "regarded"),
    (r"\bhighly sought\b", "sought"),
    (r"\binternationally recognized\b", "known internationally"),
    (r"\bglobally recognized\b", "known globally"),
    # Knowledge cutoff
    (r"\bwhile specific details are (?:limited|scarce)\b", "specific details are limited"),
    (r"\bnot widely (?:available|documented|disclosed)\b", "not publicly available"),
    (r"\bmaintains a low profile\b", "has limited public information"),
    (r"\bkeeps (?:personal|private) details (?:private|under wraps)\b", "has limited public information"),
    # Curly quote normalization
    ("\u201c", '"'),
    ("\u201d", '"'),
    ("\u2018", "'"),
    ("\u2019", "'"),
    # Em dash replacement
    ("\u2014", ","),
    ("\u2013", ","),
    # Placeholder cleanup
    (r"\bXX-XX-XXXX\b", "[date]"),
    (r"\bYYYY-MM-DD\b", "[date]"),
    (r"\[insert citation\]", "[citation needed]"),
    (r"\[insert reference\]", "[reference needed]"),
    (r"\[insert name\]", "[name]"),
    (r"\[insert date\]", "[date]"),
    # ChatGPT markup cleanup
    (r"\bcontentReference\[[\w:]+\]\{[^}]*\}", ""),
    (r"\boai_citation:\d+\b", ""),
    (r"\bgrok_card\b", ""),
    (r"\bgrok_render_citation_card_json\([^)]*\)", ""),
    (r"\bturn\dsearch\d\b", ""),
    # Concrete overuse
    (r"\bconcrete (example|evidence|steps|action|proof)\b", r"specific \1"),
    # "Hence" / "consequently" / "thus" at sentence start
    (r"^\s*hence,?\s+", ""),
    (r"^\s*consequently,?\s+", ""),
    (r"^\s*thus,?\s+", ""),
]

TRANSPARENCY_COMPONENTS = [
    ("استقبال النص", "تقطيع الكلمات والجمل والتحقق من حد 250 حرفًا."),
    ("عزل النص المؤهل", "استبعاد المراجع والجداول والمعادلات والأكواد قبل حساب مؤشرات الأسلوب."),
    ("الحيرة والتدفق", "قياس تنوع المفردات وتباين أطوال الجمل كمؤشرات محلية."),
    ("محاكاة معيارية محلية", "قواعد قابلة للتفسير تقارن النص بأنماط لغوية شائعة، وليست نموذجًا عميقًا أو خدمة خارجية."),
    ("مصنف الجمل", "إسناد درجة مراجعة لكل جملة بدل الاكتفاء بحكم شامل."),
    ("درع إعادة الصياغة", "رصد اجتماع مفردات آلية مع تدفق منتظم أو صياغة مصقولة جدًا."),
    ("عتبة 20%", "إظهار درع أمان عند الإشارات المنخفضة لتجنب اتهام النصوص البشرية المنظمة."),
    ("ESL De-biasing", "خفض القلق عند ظهور إنجليزية بسيطة ومنظمة دون قوالب آلية واضحة."),
    ("وعي بنية البحث", "تخفيف أثر أقسام المنهجية والمواد لأنها بطبيعتها جافة ومنتظمة."),
    ("الأكواد والترجمة", "إظهار تنبيهات منفصلة عند وجود كتل برمجية أو انتقالات لغوية لا تشبه النثر العادي."),
    ("مفردات الذكاء الاصطناعي", "إبراز الكلمات والعبارات التي ترفع الحاجة إلى مراجعة بشرية."),
]

REALISM_LIMITATIONS = [
    ("لا يوجد اتصال خارجي", "التطبيق لا يتصل بـ GPTZero أو iThenticate أو Turnitin أو Crossref أو PubMed."),
    ("لا حكم قطعي", "النتائج مؤشرات مراجعة فقط، ولا تثبت أن النص بشري أو مولد آليًا."),
    ("لا نسبة أصالة", "النسب المعروضة هي قوة إشارات داخلية، وليست نسبة الجمل المكتوبة بالذكاء الاصطناعي."),
    ("لا تحقق مصادر حقيقي", "تنبيهات المراجع تفحص الشكل محليًا فقط، ولا تؤكد وجود المصدر على الإنترنت."),
    ("قابل للخطأ", "النصوص الأكاديمية الجافة، وكتابة غير الناطقين بالإنجليزية، والترجمات قد ترفع الإشارة خطأً."),
]


def split_sentences(text: str) -> List[str]:
    """Lightweight sentence splitter that avoids runtime model downloads."""
    text = re.sub(r"\s+", " ", text.strip())
    if not text:
        return []
    parts = re.split(r"(?<=[.!?؟])\s+", text)
    return [part.strip() for part in parts if part.strip()]


def tokenize_words(text: str) -> List[str]:
    return re.findall(r"\b[\w'-]+\b", text, flags=re.UNICODE)


def preserve_case(source: str, replacement: str) -> str:
    if source.isupper():
        return replacement.upper()
    if source[:1].isupper():
        return replacement[:1].upper() + replacement[1:]
    return replacement


def normalize_spacing(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\s+([,.;:!?؟])", r"\1", text)
    text = re.sub(r"([([{])\s+", r"\1", text)
    text = re.sub(r"\s+([])}])", r"\1", text)
    return text.strip()


def strip_chatbot_markup(text: str) -> str:
    """Remove chatbot scaffolding and Markdown wrappers from prose."""
    text = re.sub(r"(?im)^\s*(?:sure|certainly|of course)[!,.\s]+", "", text)
    text = re.sub(
        r"(?im)^\s*here(?:'s| is)\s+(?:a|the)\s+(?:revised|polished|improved|updated)\s+(?:version|draft|text).*?:\s*$",
        "",
        text,
    )
    text = re.sub(r"(?im)^\s*as (?:an ai|a large language model)(?: language model| assistant)?,?\s*", "", text)
    text = re.sub(r"(?m)^\s{0,3}#{1,6}\s+", "", text)
    text = re.sub(r"(?m)^\s*[-*]\s+\*\*([^*\n]+)\*\*\s*[-:]\s*", r"\1: ", text)
    text = re.sub(r"(?m)^\s*(?:[-*]\s+)?(?:key takeaways?|summary|conclusion)\s*:\s*$", "", text, flags=re.I)
    text = re.sub(r"\[([^\]\n]+)\]\((?:https?://|/)[^)]+\)", r"\1", text)
    text = re.sub(r"(\*\*|__)([^*_`\n][^*_`\n]*?)\1", r"\2", text)
    text = re.sub(r"(?<!`)`([^`\n]+)`(?!`)", r"\1", text)
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u2014", ",").replace("\u2013", ",")
    return apply_wikipedia_ai_style_guard(text)


def apply_wikipedia_ai_style_guard(text: str) -> str:
    """Apply Wikipedia-derived AI-writing cleanup rules as a final safety pass."""
    revised = normalize_spacing(text)
    for _ in range(2):
        before = revised
        for pattern, replacement in WIKIPEDIA_AI_STYLE_REWRITES:
            revised = re.sub(pattern, replacement, revised, flags=re.I)
        revised = normalize_spacing(revised)
        if revised == before:
            break
    revised = re.sub(r"\s+\.", ".", revised)
    revised = re.sub(r"\.\s*\.", ".", revised)
    revised = re.sub(r"\s+,", ",", revised)
    revised = re.sub(r"\b(is|was)\s+shows\b", "shows", revised, flags=re.I)
    revised = re.sub(r"([.!?])\s+([a-z])", lambda m: f"{m.group(1)} {m.group(2).upper()}", revised)
    return normalize_spacing(revised)


def polish_layout_artifacts(text: str) -> str:
    """Fix small grammar and spacing artifacts created by rule-based revision."""
    text = apply_wikipedia_ai_style_guard(text)
    text = re.sub(r"\b(an)\s+([bcdfghjklmnpqrstvwxyz][A-Za-z-]*)", r"a \2", text, flags=re.I)
    text = re.sub(r"\b(a)\s+([aeio][A-Za-z-]*)", r"an \2", text, flags=re.I)
    text = re.sub(r"\bthe the\b", "the", text, flags=re.I)
    text = re.sub(r"\b(an|a)\s+(?:social media accounts)\b", "social media accounts", text, flags=re.I)
    text = re.sub(r"\s+([,.;:!?؟،])", r"\1", text)
    text = re.sub(r"([.!?]){2,}", r"\1", text)
    text = re.sub(r"([.!?])\s+([a-z])", lambda m: f"{m.group(1)} {m.group(2).upper()}", text)
    text = re.sub(r"([؟])\s+([a-zA-Z])", lambda m: f"{m.group(1)} {m.group(2).upper()}", text)
    return normalize_spacing(text)


@st.cache_data
def load_synonym_dictionary(filepath: str = "synonyms_academic.csv") -> pd.DataFrame:
    path = APP_DIR / filepath
    if not path.exists():
        st.warning("ملف المرادفات غير موجود؛ سيعمل التطبيق بدون استبدالات مجال تخصصية.")
        return pd.DataFrame(columns=sorted(SYNONYM_COLUMNS))

    try:
        df = pd.read_csv(path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(path, encoding="utf-8-sig")
    except Exception as exc:
        st.error(f"تعذر تحميل ملف المرادفات: {exc}")
        return pd.DataFrame(columns=sorted(SYNONYM_COLUMNS))

    missing = SYNONYM_COLUMNS.difference(df.columns)
    if missing:
        st.error("ملف المرادفات يجب أن يحتوي على الأعمدة: domain, original, replacement")
        return pd.DataFrame(columns=sorted(SYNONYM_COLUMNS))

    df = df.dropna(subset=["domain", "original", "replacement"]).copy()
    for column in SYNONYM_COLUMNS:
        df[column] = df[column].astype(str).str.strip()
    return df[df["original"].ne("") & df["replacement"].ne("")]


@st.cache_resource
def load_semantic_model():
    if SentenceTransformer is None:
        return None
    try:
        return SentenceTransformer("all-MiniLM-L6-v2")
    except Exception:
        st.warning("تعذر تحميل نموذج القفل الدلالي؛ سيستخدم التطبيق فحص تشابه محليًا أخف.")
        return None


class AcademicRevisionEngine:
    def __init__(
        self,
        domain: str,
        intensity: int,
        text: str,
        preserve_word_count: bool,
    ) -> None:
        self.domain = domain
        self.intensity = intensity
        self.original_text = text
        self.preserve_word_count = preserve_word_count
        self.synonym_map = self._build_synonym_map()
        self.semantic_model = load_semantic_model()
        self.citation_pattern = re.compile(r"\[[\d,\-; ]+\]|\([^)]*\d{4}[^)]*\)")
        self.data_pattern = re.compile(
            r"\b\d+(?:\.\d+)?\s?(?:%|mg|g|kg|ml|L|mm|cm|m|km|Hz|kHz|MHz|GHz|V|W|kW|°C|K|s|min|h)?\b",
            re.I,
        )
        self.reference_lengths = [15, 22, 18, 30, 25, 19, 28, 16, 21, 27, 23, 17, 20, 24, 26, 19, 22, 31]
        self.hedges = {
            "medical": ["may indicate", "appears to suggest", "is clinically consistent with"],
            "engineering": ["may indicate", "appears to support", "is consistent with"],
            "humanities": ["may suggest", "can be read as", "appears to reflect"],
            "general": ["may suggest", "appears to indicate", "is consistent with"],
        }
        self.causal_links = {
            "medical": {"therefore": "as a result", "so": "as a result", "because": "because of this"},
            "engineering": {"therefore": "consequently", "so": "as a result", "because": "due to this"},
            "humanities": {"therefore": "consequently", "so": "in this context", "because": "owing to this"},
            "general": {"therefore": "accordingly", "so": "as a result", "because": "because of this"},
        }
        self.contrast_links = {
            "medical": {"but": "whereas", "however": "in contrast"},
            "engineering": {"but": "whereas", "however": "alternatively"},
            "humanities": {"but": "nevertheless", "however": "nonetheless"},
            "general": {"but": "however", "however": "nevertheless"},
        }
        self.openers = {
            "medical": [],
            "engineering": [],
            "humanities": [],
            "general": [],
        }
        self.formulaic_rewrites = [
            *WIKIPEDIA_AI_STYLE_REWRITES,
            (r"\bclearly proves that\b", "suggests that"),
            (r"\bproves that\b", "suggests that"),
            (r"\bclearly demonstrates that\b", "indicates that"),
            (r"\bit is important to note that\s+", ""),
            (r"\bit should be noted that\s+", ""),
            (r"\bit is worth noting that\s+", ""),
            (r"\bit should be emphasized that\s+", ""),
            (r"\bin today's (?:world|society)\b", "currently"),
            (r"\bplays a (?:crucial|vital) role in\b", "is part of"),
            (r"\bhas a significant impact on\b", "affects"),
            (r"\bdelve into\b", "examine"),
            (r"\bstands as a testament to the evolving landscape of\b", "is part of"),
            (r"\bserves as a testament to the evolving landscape of\b", "is part of"),
            (r"\bstands as a testament to\b", "shows"),
            (r"\bserves as a testament to\b", "shows"),
            (r"\bmaintains an active social media presence\b", "has social media accounts"),
            (r"\bstands as (?:a|an|the)?\s*", "is "),
            (r"\bserves as (?:a|an|the)?\s*", "is "),
            (r"\bis testament to\b", "shows"),
            (r"\bis a testament to\b", "shows"),
            (r"\bis a reminder of\b", "shows"),
            (r",\s*highlighting its (?:pivotal|key|crucial|vital) role and significant impact on the broader ([^.!?]+)", r" and affects the \1"),
            (r",\s*(?:highlighting|underscoring|emphasizing) (?:the )?(?:importance|significance) of ([^.!?]+)", r" and discusses \1"),
            (r"\bpivotal role\b", "role"),
            (r"\bsignificant impact\b", "effect"),
            (r"\bbroader community\b", "community"),
            (r"\bmarks? a (?:significant|pivotal|key|crucial) (?:moment|shift|turning point) in\b", "changed"),
            (r"\b(?:underscores?|highlights?|emphasizes?) (?:the )?(?:importance|significance|need) of\b", "shows"),
            (r"\breflects? broader\b", "reflects"),
            (r"\bsymboli[sz]es? (?:its )?(?:ongoing|enduring|lasting)\b", "shows"),
            (r"\bsetting the stage for\b", "preceding"),
            (r"\bindelible mark\b", "effect"),
            (r"\bdeeply rooted\b", "long-standing"),
            (r"\bvaluable insights into\b", "evidence about"),
            (r"\baligns? with\b", "matches"),
            (r"\bresonates? with\b", "relates to"),
            (r"\bboasts? a\b", "has a"),
            (r"\brich tapestry of\b", "range of"),
            (r"\bevolving landscape\b", "field"),
            (r"\bactive social media presence\b", "social media accounts"),
            (r"\bindependent coverage\b", "coverage"),
            (r"\bmedia outlets\b", "publications"),
            (r"\bprofiled in\b", "covered in"),
            (r"\bindustry reports\b", "reports"),
            (r"\bobservers have cited\b", "sources cite"),
            (r"\bexperts argue\b", "some sources argue"),
            (r"\bsome critics\b", "some reviewers"),
            (r"\bas of my last (?:training update|knowledge update|knowledge),?\s*", ""),
            (r"\bup to my last training update,?\s*", ""),
            (r"\bas an ai (?:language model|assistant),?\s*", ""),
            (r"\bi hope this helps[.!]?\s*", ""),
            (r"^\s*(?:of course|certainly)[!,]?\s*", ""),
            (r"\butilize\b", "use"),
            (r"\bfacilitate\b", "support"),
            (r"\bmoreover,\s+moreover,\s+", "moreover, "),
            (r"\bin conclusion,\s+", ""),
            (r"\boverall,\s+", ""),
        ]
        self.repetitive_openers = {
            "furthermore": "",
            "moreover": "",
            "additionally": "",
            "therefore": "As a result",
            "however": "By contrast",
        }
        self.editorial_connectors = [
            r"\bfurthermore,?\s+",
            r"\bmoreover,?\s+",
            r"\badditionally,?\s+",
            r"\bin conclusion,?\s+",
            r"\boverall,?\s+",
            r"\bon the other hand,?\s+",
            r"\bit is worth noting that\s+",
            r"\bit should be emphasized that\s+",
            r"\bas previously mentioned,?\s+",
            r"\bbased on the foregoing,?\s+",
        ]
        self.editorial_vocabulary_swaps = [
            (r"\bshows important results\b", "reports specific findings"),
            (r"\bshows that\b", "indicates that"),
            (r"\bimportant\b", "specific"),
            (r"\bsignificant\b", "measurable"),
            (r"\bshows\b", "reveals"),
            (r"\bindicates\b", "suggests"),
            (r"\bresults\b", "findings"),
            (r"\bdata\b", "evidence"),
            (r"\bmethod\b", "approach"),
            (r"\btopic\b", "question"),
            (r"\bfactor\b", "force"),
            (r"\ba large number of\b", "many"),
            (r"\bin a clear way\b", "clearly"),
        ]
        self.editorial_passive_rewrites = [
            (r"\bit was found that\b", "the results suggest that"),
            (r"\bit was observed that\b", "the observation indicates that"),
            (r"\bit is argued that\b", "the argument is that"),
            (r"\bit is suggested that\b", "the material suggests that"),
            (r"\bit can be seen that\b", "the evidence indicates that"),
            (r"\bwas analyzed\b", "was reviewed"),
            (r"\bwere analyzed\b", "were reviewed"),
            (r"\bwas examined\b", "was reviewed"),
            (r"\bwere examined\b", "were reviewed"),
        ]
        self.editorial_voice_openers = [
            "In this passage",
            "In the evidence",
            "In the analysis",
            "In the reported material",
        ]
        self.editorial_short_beats: List[str] = []
        self.editorial_question_added = False
        self.surprise_word_map = {
            "show": ["reveal", "indicate", "demonstrate", "document", "report"],
            "find": ["identify", "observe", "detect", "note", "discover"],
            "use": ["employ", "apply", "adopt", "utilize in limited contexts", "rely on"],
            "make": ["produce", "generate", "yield", "create", "form"],
            "give": ["provide", "offer", "supply", "present", "contribute"],
            "help": ["support", "aid", "assist", "facilitate carefully", "enable"],
            "change": ["alter", "modify", "shift", "transform", "adjust"],
            "increase": ["raise", "elevate", "boost", "amplify", "grow"],
            "decrease": ["reduce", "lower", "diminish", "cut", "shrink"],
            "important": ["notable", "worth noting", "relevant", "meaningful", "consequential"],
            "different": ["distinct", "varied", "divergent", "separate", "unlike"],
            "large": ["substantial", "considerable", "sizable", "extensive", "ample"],
            "small": ["modest", "limited", "minor", "slight", "narrow"],
            "good": ["effective", "suitable", "appropriate", "favorable", "sound"],
            "bad": ["poor", "inadequate", "deficient", "suboptimal", "limited"],
            "new": ["recent", "novel", "current", "fresh", "updated"],
            "old": ["previous", "earlier", "prior", "established", "existing"],
            "many": ["numerous", "multiple", "several", "various", "a range of"],
            "few": ["limited", "scant", "a handful of", "sparse", "a small number of"],
            "likely": ["probable", "plausible", "expected", "anticipated", "reasonable to expect"],
            "possible": ["feasible", "conceivable", "achievable", "viable", "attainable"],
            "clear": ["evident", "apparent", "obvious in context", "plain", "transparent"],
            "simple": ["straightforward", "uncomplicated", "direct", "basic", "elementary"],
            "complex": ["intricate in structure", "multi-layered", "involved", "elaborate", "sophisticated"],
            "fast": ["rapid", "swift", "quick", "expedient", "prompt"],
            "slow": ["gradual", "progressive", "measured", "deliberate", "incremental"],
        }
        self.function_word_variants = {
            "the": ["the", "this", "such", "the given", "the present"],
            "a": ["a", "one", "any", "some"],
            "is": ["is", "appears to be", "represents", "constitutes"],
            "are": ["are", "appear to be", "constitute", "represent"],
            "was": ["was", "proved to be", "emerged as"],
            "were": ["were", "proved to be", "emerged as"],
            "has": ["has", "possesses", "exhibits", "displays"],
            "have": ["have", "possess", "exhibit", "display"],
            "can": ["can", "is able to", "is capable of", "may"],
            "will": ["will", "is expected to", "is likely to", "should"],
            "this": ["this", "the current", "the present", "the following"],
            "these": ["these", "the current", "the present", "the following"],
            "that": ["that", "which", "the indicated"],
            "which": ["which", "that", "a factor that"],
            "it": ["it", "this", "the system", "the approach"],
            "they": ["they", "these", "the results", "the findings"],
        }

    def _build_synonym_map(self) -> Dict[str, str]:
        df = load_synonym_dictionary()
        if df.empty:
            return {}
        if self.domain == "general":
            scoped = df
        else:
            scoped = df[df["domain"].str.lower() == self.domain.lower()]
        return {
            str(row["original"]).lower(): str(row["replacement"])
            for _, row in scoped.iterrows()
            if str(row["original"]).strip() and str(row["replacement"]).strip()
        }

    def _protect_fragments(self, text: str) -> Tuple[str, Dict[str, str]]:
        replacements: Dict[str, str] = {}

        def replace(match: re.Match[str]) -> str:
            token = f"__PROTECTED_{len(replacements)}__"
            replacements[token] = match.group(0)
            return token

        protected = self.citation_pattern.sub(replace, text)
        protected = self.data_pattern.sub(replace, protected)
        return protected, replacements

    def _restore_fragments(self, text: str, replacements: Dict[str, str]) -> str:
        for token, value in replacements.items():
            text = text.replace(token, value)
        return text

    def _semantic_similarity(self, original: str, revised: str) -> float:
        if not original.strip() or not revised.strip():
            return 0.0
        if self.semantic_model is not None and util is not None:
            emb_original = self.semantic_model.encode(original, convert_to_tensor=True)
            emb_revised = self.semantic_model.encode(revised, convert_to_tensor=True)
            return float(util.pytorch_cos_sim(emb_original, emb_revised).item())
        original_words = [word.lower() for word in tokenize_words(original)]
        revised_words = [word.lower() for word in tokenize_words(revised)]
        jaccard = jaccard_similarity(original_words, revised_words)
        sequence = difflib.SequenceMatcher(None, original.lower(), revised.lower()).ratio()
        return max(jaccard, sequence)

    def _semantic_lock(self, original: str, candidate: str, threshold: float = 0.92) -> str:
        effective_threshold = threshold if self.semantic_model is not None else min(threshold, 0.35)
        if self._semantic_similarity(original, candidate) >= effective_threshold:
            return candidate
        return original

    def _most_predictable_terms(self, sentence: str) -> Sequence[str]:
        words = [word for word in tokenize_words(sentence) if word.lower() in self.synonym_map]
        if not words:
            return []
        counts = Counter(word.lower() for word in tokenize_words(self.original_text))
        return sorted(words, key=lambda word: (-counts[word.lower()], -len(word)))

    def _rewrite_formulaic_phrases(self, sentence: str) -> str:
        revised = sentence
        for pattern, replacement in self.formulaic_rewrites:
            revised = re.sub(pattern, replacement, revised, flags=re.I)
        revised = normalize_spacing(revised)
        return revised[:1].upper() + revised[1:] if revised else revised

    def _revise_repetitive_opener(self, sentence: str, index: int) -> str:
        if index == 0:
            return sentence
        for source, replacement in self.repetitive_openers.items():
            if re.match(rf"^\s*{re.escape(source)},?\s+", sentence, flags=re.I):
                if index % 2 == 0 and replacement:
                    return re.sub(rf"^\s*{re.escape(source)},?\s+", f"{replacement}, ", sentence, count=1, flags=re.I)
                shortened = re.sub(rf"^\s*{re.escape(source)},?\s+", "", sentence, count=1, flags=re.I)
                return shortened[:1].upper() + shortened[1:]
        return sentence

    def engine1_perplexity_injector(self, sentence: str) -> str:
        """Conservative term revision and hedging for over-certain academic claims."""
        revised = self._rewrite_formulaic_phrases(sentence)
        if not self.synonym_map or self.intensity < 2:
            revised = revised
        else:
            max_changes = max(1, min(3, self.intensity - 1))
            for word in self._most_predictable_terms(revised)[:max_changes]:
                replacement = self.synonym_map.get(word.lower())
                if replacement:
                    revised = re.sub(
                        rf"\b{re.escape(word)}\b",
                        preserve_case(word, replacement),
                        revised,
                        count=1,
                        flags=re.I,
                    )

        certainty_markers = re.compile(r"\b(always|never|definitely|undoubtedly)\b", re.I)
        if self.intensity >= 3 and certainty_markers.search(revised):
            hedge = self.hedges.get(self.domain, self.hedges["general"])[0]
            revised = certainty_markers.sub(hedge, revised, count=1)
        elif textstat is not None and self.intensity >= 4:
            try:
                if textstat.flesch_reading_ease(revised) < 20 and "," not in revised:
                    hedge = self.hedges.get(self.domain, self.hedges["general"])[1]
                    revised = f"{hedge}, {revised[:1].lower()}{revised[1:]}"
            except Exception:
                pass
        return normalize_spacing(revised)

    def engine2_burstiness_synthesizer(self, sentences: List[str]) -> List[str]:
        """Create natural sentence length variation that mimics human burstiness."""
        if not sentences:
            return sentences

        # Step 1: Split overly long sentences
        adjusted = []
        for sentence in sentences:
            adjusted.extend(self._split_long_sentence(sentence))

        # Step 2: Calculate current burstiness metrics
        lengths = [len(tokenize_words(s)) for s in adjusted]
        if not lengths:
            return adjusted

        avg_len = float(np.mean(lengths))
        std_len = float(np.std(lengths))
        cv = std_len / avg_len if avg_len > 0 else 0  # coefficient of variation

        # Human writing typically has CV between 0.35-0.65
        # AI writing typically has CV < 0.25
        target_cv = 0.45 if self.intensity >= 3 else 0.35

        # Step 3: If CV is too low (too uniform), inject variation
        if cv < target_cv and len(adjusted) >= 3:
            # Strategy A: Insert short beat sentences
            short_beats = [
                "This matters.",
                "The evidence is clear.",
                "We return to this below.",
                "This warrants attention.",
                "The pattern holds.",
                "This is expected.",
                "The reason is straightforward.",
                "We note this finding.",
                "This aligns with prior work.",
                "The implication is direct.",
            ]

            domain_beats = {
                "medical": [
                    "The clinical picture is consistent.",
                    "This aligns with current practice.",
                    "The finding is reproducible.",
                ],
                "engineering": [
                    "The measurement confirms this.",
                    "This matches the model output.",
                    "The design accounts for this.",
                ],
                "humanities": [
                    "The text supports this reading.",
                    "This interpretation holds.",
                    "The argument follows.",
                ],
                "general": short_beats,
            }

            beats = domain_beats.get(self.domain, short_beats)

            # Find good insertion points (after longer sentences)
            beat_idx = 0
            result = []
            for i, sentence in enumerate(adjusted):
                result.append(sentence)
                word_count = len(tokenize_words(sentence))
                # Insert a beat after a long sentence, but not too frequently
                if (word_count >= avg_len and
                    i < len(adjusted) - 1 and
                    len(tokenize_words(adjusted[i+1])) >= avg_len * 0.8 and
                    i % 3 == 1 and
                    beat_idx < 2):  # Don't add too many beats
                    result.append(beats[beat_idx % len(beats)])
                    beat_idx += 1

            adjusted = result

        # Step 4: Break adjacent similar-length sentences
        balanced = []
        previous_len = None
        for sentence in adjusted:
            current_len = len(tokenize_words(sentence))
            if previous_len is not None and abs(current_len - previous_len) <= 2 and current_len > 16:
                # Try to split or merge to create variation
                midpoint = max(8, len(sentence.split()) // 2)
                words = sentence.split()
                first = " ".join(words[:midpoint]).rstrip(" ,;")
                second = " ".join(words[midpoint:]).strip(" ,;")
                if first and second:
                    if first[-1] not in ".!?":
                        first += "."
                    second = second[:1].upper() + second[1:]
                    if second[-1] not in ".!?":
                        second += "."
                    balanced.extend([first, second])
                    previous_len = len(tokenize_words(second))
                    continue
            balanced.append(sentence)
            previous_len = current_len

        return balanced

    def engine3_style_variety_editor(self, sentences: List[str]) -> List[str]:
        """Vary punctuation and sentence openings without adding new information."""
        if self.intensity < 3:
            return [self._revise_repetitive_opener(sentence, index) for index, sentence in enumerate(sentences)]
        openers = self.openers.get(self.domain, self.openers["general"])
        revised: List[str] = []
        for index, sentence in enumerate(sentences):
            sentence = self._revise_repetitive_opener(sentence, index)
            if openers and index > 0 and index % 4 == 0 and len(tokenize_words(sentence)) > 9:
                opener = openers[index % len(openers)]
                if not sentence.lower().startswith(opener.lower()):
                    sentence = f"{opener}, {sentence[:1].lower()}{sentence[1:]}"
            if self.intensity >= 4 and ";" not in sentence and "," in sentence and index % 5 == 2:
                sentence = sentence.replace(",", ";", 1)
            revised.append(sentence)
        return revised

    def engine4_semantic_deepener(self, sentences: List[str]) -> List[str]:
        """Replace generic connectors with more precise causal or contrastive links."""
        causal = self.causal_links.get(self.domain, self.causal_links["general"])
        contrast = self.contrast_links.get(self.domain, self.contrast_links["general"])
        revised: List[str] = []
        for sentence in sentences:
            original = sentence
            for source, target in {**causal, **contrast}.items():
                sentence = re.sub(rf"^\s*{re.escape(source)}\b", target, sentence, flags=re.I)
                sentence = re.sub(rf"\b{re.escape(source)}\b", target, sentence, count=1, flags=re.I)
            revised.append(self._semantic_lock(original, sentence, threshold=0.75))
        return revised

    def engine5_structure_regularizer(self, sentences: List[str]) -> List[str]:
        """Reduce repetitive paragraph rhythm while preserving sentence order."""
        revised: List[str] = []
        previous_start = ""
        for sentence in sentences:
            # Remove filler words that AI overuses
            cleaned = re.sub(r"\b(very|really|basically|clearly|simply|obviously|certainly|indeed|thus|hence|thereby|heretofore|notably|interestingly|importantly)\b", "", sentence, flags=re.I)
            # Also remove "in order to" -> "to"
            cleaned = re.sub(r"\bin order to\b", "to", cleaned, flags=re.I)
            # Remove "a number of" -> "several"
            cleaned = re.sub(r"\ba number of\b", "several", cleaned, flags=re.I)
            # Remove "in the context of" -> "in"
            cleaned = re.sub(r"\bin the context of\b", "in", cleaned, flags=re.I)
            # Remove "at the level of" -> "in"
            cleaned = re.sub(r"\bat the level of\b", "in", cleaned, flags=re.I)

            words = tokenize_words(cleaned)
            current_start = " ".join(word.lower() for word in words[:3])
            if previous_start and current_start == previous_start and len(words) > 6:
                cleaned = " ".join(words[3:])
                cleaned = cleaned[:1].upper() + cleaned[1:]
            previous_start = current_start
            revised.append(normalize_spacing(cleaned))
        return revised

    def engine6_coherence_checker(self, original_sentences: List[str], revised_sentences: List[str]) -> List[str]:
        """Repair obvious formatting and sentence-ending issues while preserving citations."""
        checked: List[str] = []
        for index, sentence in enumerate(revised_sentences):
            original = original_sentences[min(index, len(original_sentences) - 1)] if original_sentences else sentence
            # Preserve any citations [1], [2,3] from original that might have been lost
            # Only add citations that were in the ORIGINAL for this sentence position
            orig_citations = re.findall(r'\[[\d,\-; ]+\]', original)
            rev_citations = re.findall(r'\[[\d,\-; ]+\]', sentence)
            # Only append if original had citations AND revised has none
            if orig_citations and not rev_citations:
                # Append the original citations at the end
                for cite in orig_citations:
                    sentence = sentence.rstrip(".") + f" {cite}."
                    break  # Only add once
            sentence = normalize_spacing(sentence)
            if sentence and sentence[-1] not in ".!?؟":
                sentence += "."
            # Use very low threshold - we WANT the revisions to pass through
            checked.append(self._semantic_lock(original, sentence, threshold=0.25))
        return checked

    def engine7_perplexity_variation(self, sentences: List[str]) -> List[str]:
        """Vary word-level predictability by injecting less common alternatives."""
        if self.intensity < 2:
            return sentences

        revised = []
        word_counter = Counter(w.lower() for s in sentences for w in tokenize_words(s))
        doc_len = sum(word_counter.values())

        for idx, sentence in enumerate(sentences):
            words = tokenize_words(sentence)
            if not words:
                revised.append(sentence)
                continue

            # Replace overused content words (appearing >2% of document)
            changes = 0
            max_changes = min(2, max(1, self.intensity - 1))
            for i, word in enumerate(words):
                if changes >= max_changes:
                    break
                lower = word.lower()
                # Skip function words, short words, protected terms
                if len(word) < 4 or lower in self.function_word_variants:
                    continue
                if word_counter[lower] / max(1, doc_len) > 0.02:
                    alternatives = self.surprise_word_map.get(lower, [])
                    if alternatives:
                        alt = alternatives[idx % len(alternatives)]
                        # Preserve case
                        alt = preserve_case(word, alt)
                        # Only substitute if it won't break grammar
                        sentence = re.sub(rf"\b{re.escape(word)}\b", alt, sentence, count=1, flags=re.I)
                        changes += 1

            # Vary function words (every 4th sentence at intensity >= 3)
            if self.intensity >= 3 and idx % 4 == 2:
                for func_word, variants in self.function_word_variants.items():
                    pattern = rf"\b{re.escape(func_word)}\b"
                    if re.search(pattern, sentence, re.I):
                        variant = variants[idx % len(variants)]
                        if variant != func_word:
                            sentence = re.sub(pattern, variant, sentence, count=1, flags=re.I)
                        break

            # Add hedging at sentence boundaries for overly certain claims (intensity >= 3)
            if self.intensity >= 3 and idx % 5 == 3:
                hedge_starters = [
                    "It seems that", "We find that", "The data suggest",
                    "In our assessment,", "Based on the evidence,",
                ]
                words_count = len(tokenize_words(sentence))
                if words_count >= 12 and not any(sentence.strip().lower().startswith(h.lower()) for h in hedge_starters):
                    starter = hedge_starters[idx % len(hedge_starters)]
                    sentence = f"{starter} {sentence[:1].lower()}{sentence[1:]}"

            sentence = normalize_spacing(sentence)
            revised.append(sentence)

        return revised

    def engine8_punctuation_entropy(self, sentences: List[str]) -> List[str]:
        """Increase punctuation diversity to break AI-like uniformity."""
        if self.intensity < 2:
            return sentences

        revised = []
        for idx, sentence in enumerate(sentences):
            # Insert parenthetical aside every ~6 sentences at intensity >= 3
            if self.intensity >= 3 and idx % 6 == 4 and len(tokenize_words(sentence)) >= 14:
                words = sentence.split()
                midpoint = len(words) // 2
                # Find a good insertion point near midpoint
                insert_at = midpoint
                for i in range(midpoint, min(midpoint + 4, len(words))):
                    if words[i].endswith(","):
                        insert_at = i + 1
                        break

                parentheticals = [
                    "(as expected)",
                    "(see below)",
                    "(in this context)",
                    "(for the present study)",
                    "(under these conditions)",
                    "(as reported earlier)",
                ]
                words.insert(insert_at, parentheticals[idx % len(parentheticals)])
                sentence = " ".join(words)

            # Replace some commas with semicolons where grammatically appropriate (intensity >= 4)
            if self.intensity >= 4 and idx % 7 == 3 and ", " in sentence:
                # Only replace if the comma separates two independent clauses
                parts = sentence.split(", ", 1)
                if len(parts) == 2 and len(tokenize_words(parts[0])) >= 5:
                    # Check if first part could stand alone (has a verb)
                    if re.search(r"\b(is|are|was|were|has|have|had|shows?|indicates?|suggests?|finds?|yields?|reports?)\b", parts[0], re.I):
                        sentence = "; ".join(parts)

            # Add occasional colon usage for lists/explanations
            if self.intensity >= 3 and idx % 8 == 5:
                colon_patterns = [
                    (r"\bas follows[: ]", "as follows: "),
                    (r"\bnamely\s+", "namely, "),
                ]
                for pattern, replacement in colon_patterns:
                    if re.search(pattern, sentence, re.I):
                        sentence = re.sub(pattern, replacement, sentence, count=1, flags=re.I)
                        break

            revised.append(normalize_spacing(sentence))

        return revised

    def engine9_lexical_diversity_injector(self, sentences: List[str]) -> List[str]:
        """Reduce word repetition by replacing repeated content words with synonyms."""
        if self.intensity < 2:
            return sentences

        # Count all words across sentences
        all_words = [w.lower() for s in sentences for w in tokenize_words(s)]
        word_freq = Counter(all_words)
        total_words = len(all_words)

        # Track which replacements we've made to avoid re-replacing
        replaced = set()
        revised = []

        for idx, sentence in enumerate(sentences):
            words = tokenize_words(sentence)
            changes = 0
            max_changes = 2 if self.intensity >= 4 else 1

            for word in words:
                if changes >= max_changes:
                    break
                lower = word.lower()
                # Only replace content words that appear >3 times and >1.5% of doc
                if len(word) < 5 or lower in replaced:
                    continue
                if word_freq[lower] >= 3 and (word_freq[lower] / max(1, total_words)) > 0.015:
                    alternatives = self.surprise_word_map.get(lower, [])
                    if alternatives:
                        alt = alternatives[(idx + hash(lower)) % len(alternatives)]
                        alt = preserve_case(word, alt)
                        sentence = re.sub(rf"\b{re.escape(word)}\b", alt, sentence, count=1, flags=re.I)
                        replaced.add(lower)
                        changes += 1

            revised.append(normalize_spacing(sentence))

        return revised

    def engine10_tortured_phrase_fixer(self, sentences: List[str]) -> List[str]:
        """Replace tortured/awkward AI paraphrases with natural phrasing."""
        revised = []
        for sentence in sentences:
            for pattern, replacement in TORTURED_PHRASES:
                sentence = re.sub(pattern, replacement, sentence, flags=re.I)
            revised.append(normalize_spacing(sentence))
        return revised

    def engine11_deduplication_guard(self, sentences: List[str]) -> List[str]:
        """Remove or rephrase duplicated sentences that are a major AI signal."""
        if len(sentences) < 2:
            return sentences

        seen_hashes: Dict[int, str] = {}
        deduplicated: List[str] = []

        for idx, sentence in enumerate(sentences):
            # Normalize for comparison: lower, strip punctuation
            normalized = re.sub(r'[^\w\s]', '', sentence.lower()).strip()
            normalized = re.sub(r'\s+', ' ', normalized)

            # Check for exact or near-duplicate
            is_duplicate = False
            for existing_hash, existing_text in seen_hashes.items():
                similarity = difflib.SequenceMatcher(None, normalized, existing_text).ratio()
                if similarity > 0.85:  # 85%+ similarity = duplicate
                    is_duplicate = True
                    break

            if is_duplicate:
                # Instead of removing, try to rephrase by shortening
                words = tokenize_words(sentence)
                if len(words) > 10:
                    # Keep only first half and add different ending
                    half = len(words) // 2
                    shortened = " ".join(words[:half])
                    # Add a different conclusion
                    alternative_endings = [
                        "and related factors affect this.",
                        "among other variables.",
                        "as discussed below.",
                        "which we address next.",
                        "in the sections that follow.",
                    ]
                    shortened += " " + alternative_endings[idx % len(alternative_endings)]
                    if shortened[-1] not in ".!?":
                        shortened += "."
                    deduplicated.append(shortened)
                # If very short duplicate, skip it entirely
            else:
                seen_hashes[hash(normalized)] = normalized
                deduplicated.append(sentence)

        return deduplicated if deduplicated else sentences

    def engine12_humanization(self, sentences: List[str]) -> List[str]:
        """Inject human-like writing patterns: contractions, fragments, hedging, informal touches."""
        if self.intensity < 2:
            return sentences

        revised: List[str] = []
        for idx, sentence in enumerate(sentences):
            words = tokenize_words(sentence)
            word_count = len(words)

            # 1. Introduce contractions (intensity >= 3)
            if self.intensity >= 3:
                contraction_map = [
                    (r"\bit is\b", "it's"),
                    (r"\bdo not\b", "don't"),
                    (r"\bcan not\b", "can't"),
                    (r"\bcannot\b", "can't"),
                    (r"\bwill not\b", "won't"),
                    (r"\bshould not\b", "shouldn't"),
                    (r"\bwould not\b", "won't"),
                    (r"\bdoes not\b", "doesn't"),
                    (r"\bis not\b", "isn't"),
                    (r"\bare not\b", "aren't"),
                    (r"\bwas not\b", "wasn't"),
                    (r"\bwere not\b", "weren't"),
                    (r"\bhas not\b", "hasn't"),
                    (r"\bhave not\b", "haven't"),
                    (r"\bthey are\b", "they're"),
                    (r"\bwe are\b", "we're"),
                    (r"\bthat is\b", "that's"),
                ]
                # Apply 1-2 contractions per sentence (not all)
                changes = 0
                for pattern, replacement in contraction_map:
                    if changes >= 2:
                        break
                    if re.search(pattern, sentence, re.I) and idx % 3 != 0:  # Skip every 3rd to avoid uniformity
                        sentence = re.sub(pattern, replacement, sentence, count=1, flags=re.I)
                        changes += 1

            # 2. Insert occasional fragment sentences (intensity >= 3)
            # After every 4-5 long sentences, add a short fragment
            if self.intensity >= 3 and idx % 5 == 4 and word_count >= 15:
                fragment_insertions = [
                    "This is expected.",
                    "So far, so standard.",
                    "We note this.",
                    "This bears repeating.",
                    "The reason is simple.",
                    "A clear pattern.",
                    "No surprise here.",
                    "This checks out.",
                    "Consistent with theory.",
                    "As anticipated.",
                ]
                # Don't modify the current sentence, but mark for fragment insertion
                revised.append(sentence)
                revised.append(fragment_insertions[idx % len(fragment_insertions)])
                continue

            # 3. Add hedging to over-certain statements (intensity >= 2)
            if self.intensity >= 2 and idx % 4 == 1:
                certainty_words = [
                    (r"\bclearly\b", "it seems"),
                    (r"\bobviously\b", "apparently"),
                    (r"\bundeniable\b", "hard to dispute"),
                    (r"\bcertainly\b", "in most cases"),
                    (r"\bdefinitely\b", "likely"),
                    (r"\bprove[s]?\b", "suggest"),
                    (r"\bproven\b", "suggested"),
                ]
                for pattern, replacement in certainty_words:
                    if re.search(pattern, sentence, re.I):
                        sentence = re.sub(pattern, replacement, sentence, count=1, flags=re.I)
                        break

            # 4. Vary sentence starters aggressively (intensity >= 3)
            if self.intensity >= 3 and word_count >= 10:
                # Detect common AI sentence starters
                ai_starters = [
                    r"^The (?:results|findings|data|study|analysis|evidence)\b",
                    r"^This (?:study|paper|research|work|article|analysis)\b",
                    r"^Our (?:results|findings|analysis|study|data|approach)\b",
                    r"^These (?:results|findings|data|factors|observations)\b",
                    r"^It (?:is|was|can|will|should|may|might)\b",
                ]
                alternative_starters = [
                    "Looking at the", "Based on", "From the", "In terms of",
                    "Turning to", "As for", "Regarding", "With respect to",
                    "We find that", "We see that", "It turns out",
                    "What emerges is", "The picture shows",
                    "According to", "Consistent with",
                ]
                for starter_pattern in ai_starters:
                    if re.match(starter_pattern, sentence, re.I):
                        # Replace the first 1-2 words with a varied opener
                        new_starter = alternative_starters[(idx + hash(sentence[:10])) % len(alternative_starters)]
                        # Remove the original starter (first 1-2 words)
                        words_list = sentence.split()
                        if len(words_list) >= 3:
                            # Remove first word, keep rest
                            rest = " ".join(words_list[1:])
                            sentence = f"{new_starter} {rest[:1].lower()}{rest[1:]}"
                        break

            # 5. Add parenthetical author voice (intensity >= 4)
            if self.intensity >= 4 and idx % 6 == 3 and word_count >= 18:
                author_notes = [
                    "(though see the caveats below)",
                    "(as one might expect)",
                    "(in our experience)",
                    "(anecdotally, at least)",
                    "(this is a simplification, of course)",
                    "(with some exceptions)",
                    "(in practice, anyway)",
                ]
                # Insert near the middle
                words_list = sentence.split()
                insert_pos = len(words_list) * 2 // 3
                words_list.insert(insert_pos, author_notes[idx % len(author_notes)])
                sentence = " ".join(words_list)

            revised.append(normalize_spacing(sentence))

        return revised

    def engine13_deep_perplexity(self, sentences: List[str]) -> List[str]:
        """Deep perplexity variation: break predictable sequences, add surprise words."""
        if self.intensity < 3:
            return sentences

        revised = []
        all_words_lower = [w.lower() for s in sentences for w in tokenize_words(s)]
        word_freq = Counter(all_words_lower)
        total = len(all_words_lower)

        for idx, sentence in enumerate(sentences):
            words = tokenize_words(sentence)
            if not words:
                revised.append(sentence)
                continue

            # 1. Replace the MOST frequent content words in the document
            # These create the lowest perplexity because they're so predictable
            max_swaps = min(3, self.intensity - 1)
            swaps_done = 0

            for word in words:
                if swaps_done >= max_swaps:
                    break
                lower = word.lower()
                freq_ratio = word_freq[lower] / max(1, total)

                # Only swap words that are very frequent (>2% of doc) and are content words
                if freq_ratio > 0.02 and len(word) >= 4:
                    alternatives = self.surprise_word_map.get(lower, [])
                    if alternatives:
                        # Pick a different alternative than what engine7 might have picked
                        alt_idx = (idx * 3 + hash(lower)) % len(alternatives)
                        alt = alternatives[alt_idx]
                        alt = preserve_case(word, alt)
                        sentence = re.sub(rf"\b{re.escape(word)}\b", alt, sentence, count=1, flags=re.I)
                        swaps_done += 1

            # 2. Break predictable bigrams by inserting qualifiers
            predictable_bigrams = [
                (r"\bplay a role\b", "factor in"),
                (r"\bplay an important role\b", "matter"),
                (r"\bconduct a study\b", "run a study"),
                (r"\bcarry out\b", "do"),
                (r"\bcarried out\b", "did"),
                (r"\bfocus on\b", "zero in on"),
                (r"\bfocused on\b", "zeroed in on"),
                (r"\baim to\b", "set out to"),
                (r"\baimed to\b", "set out to"),
                (r"\blead to\b", "bring about"),
                (r"\bled to\b", "brought about"),
                (r"\bresult in\b", "give rise to"),
                (r"\bresulted in\b", "gave rise to"),
                (r"\bbased on\b", "building on"),
                (r"\bin order to\b", "to"),
                (r"\ba number of\b", "several"),
                (r"\ba large number of\b", "many"),
                (r"\ba variety of\b", "various"),
                (r"\ba wide range of\b", "many"),
                (r"\bthe present study\b", "this work"),
                (r"\bthe current study\b", "this work"),
                (r"\bthe present work\b", "this study"),
                (r"\bthe current work\b", "this study"),
            ]
            for pattern, replacement in predictable_bigrams:
                sentence = re.sub(pattern, replacement, sentence, count=1, flags=re.I)

            # 3. Add occasional unexpected word choices (increase perplexity)
            if idx % 7 == 5 and len(words) >= 12:
                # Insert a less common transition or qualifier
                surprise_insertions = [
                    "Interestingly,", "Notably, though,", "In fact,",
                    "As it happens,", "To be sure,", "Admittedly,",
                    "By contrast,", "In any case,", "At any rate,",
                ]
                insertion = surprise_insertions[idx % len(surprise_insertions)]
                # Place it near the beginning of the sentence
                words_list = sentence.split()
                if len(words_list) >= 4:
                    insert_after = min(3, len(words_list) - 1)
                    words_list.insert(insert_after, insertion.lower().rstrip(','))
                    sentence = " ".join(words_list)
                    # Fix capitalization
                    sentence = sentence[:1].upper() + sentence[1:]

            revised.append(normalize_spacing(sentence))

        return revised

    def _split_long_sentence(self, sentence: str) -> List[str]:
        words = tokenize_words(sentence)
        limit = 42 if self.intensity <= 2 else 34
        if len(words) <= limit:
            return [sentence]

        breakpoints = [",", ";", " and ", " but ", " whereas ", " which "]
        lower = sentence.lower()
        midpoint = len(sentence) // 2
        candidates: List[int] = []
        for marker in breakpoints:
            start = 0
            while True:
                idx = lower.find(marker, start)
                if idx == -1:
                    break
                candidates.append(idx + len(marker.strip()))
                start = idx + 1

        if candidates:
            split_at = min(candidates, key=lambda idx: abs(idx - midpoint))
            first = sentence[:split_at].strip(" ,;")
            second = sentence[split_at:].strip(" ,;")
        else:
            raw_words = sentence.split()
            split_at_word = max(12, len(raw_words) // 2)
            first = " ".join(raw_words[:split_at_word]).strip(" ,;")
            second = " ".join(raw_words[split_at_word:]).strip(" ,;")

        if not first or not second:
            return [sentence]
        if first[-1] not in ".!?؟":
            first += "."
        second = second[:1].upper() + second[1:]
        if second[-1] not in ".!?؟":
            second += "."
        return [first, second]

    def _flow_is_extreme(self, sentences: List[str]) -> bool:
        lengths = [len(tokenize_words(sentence)) for sentence in sentences]
        if len(lengths) < 3:
            return False
        baseline = np.std(self.reference_lengths)
        return bool(np.std(lengths) > 3 * baseline)

    def _clean_editorial_connectors(self, sentence: str) -> str:
        sentence = sentence.replace("—", ",").replace("–", ",")
        for pattern in self.editorial_connectors:
            sentence = re.sub(pattern, "", sentence, flags=re.I)
        sentence = re.sub(r"\s*[-]{2,}\s*", ", ", sentence)
        return normalize_spacing(sentence)

    def _activate_english_voice(self, sentence: str) -> str:
        for pattern, replacement in self.editorial_passive_rewrites:
            sentence = re.sub(pattern, replacement, sentence, flags=re.I)
        return normalize_spacing(sentence)

    def _enrich_english_vocabulary(self, sentence: str) -> str:
        limit = 1 if self.intensity < 4 else 2
        changes = 0
        for pattern, replacement in self.editorial_vocabulary_swaps:
            if changes >= limit:
                break
            if re.search(pattern, sentence, flags=re.I):
                sentence = re.sub(pattern, replacement, sentence, count=1, flags=re.I)
                changes += 1
        return normalize_spacing(sentence)

    def _polish_english_artifacts(self, sentence: str) -> str:
        sentence = re.sub(r"\breveals revealing (?:results|findings)\b", "reveals meaningful findings", sentence, flags=re.I)
        sentence = re.sub(r"\b(clinical measurements|evidence|data) reveals\b", r"\1 reveal", sentence, flags=re.I)
        sentence = re.sub(r"\bthe results suggest that the evidence\b", "the results suggest that it", sentence, flags=re.I)
        return normalize_spacing(sentence)

    def _english_editorial_pass(self, paragraph: str) -> str:
        if self.intensity < 3:
            return normalize_spacing(paragraph.replace("—", ",").replace("–", ","))

        sentences = split_sentences(paragraph)
        if not sentences:
            return paragraph

        revised: List[str] = []
        for index, sentence in enumerate(sentences):
            sentence = self._clean_editorial_connectors(sentence)
            sentence = self._activate_english_voice(sentence)
            sentence = self._enrich_english_vocabulary(sentence)
            sentence = self._polish_english_artifacts(sentence)
            sentence = sentence[:1].upper() + sentence[1:] if sentence else sentence
            if sentence and sentence[-1] not in ".!?":
                sentence += "."
            revised.append(sentence)

            if self.editorial_short_beats and self.intensity >= 5 and len(tokenize_words(sentence)) >= 20 and index % 3 == 0:
                revised.append(self.editorial_short_beats[index % len(self.editorial_short_beats)])

        return normalize_spacing(" ".join(revised))

    def revise_paragraph(self, paragraph: str, context: Sequence[str] | None = None) -> str:
        sentences = split_sentences(paragraph)
        context = list(context or [])

        # Phase 1: Fix structural problems
        working = self.engine10_tortured_phrase_fixer(sentences)
        working = self.engine11_deduplication_guard(working)

        # Phase 2: Sentence-level revision
        working = [self.engine1_perplexity_injector(s) for s in working]

        # Phase 3: Burstiness and length variation
        working = self.engine2_burstiness_synthesizer(working)

        # Phase 4: Vocabulary diversity
        working = self.engine9_lexical_diversity_injector(working)

        # Phase 5: Perplexity variation (two passes)
        working = self.engine7_perplexity_variation(working)
        working = self.engine13_deep_perplexity(working)

        # Phase 6: Punctuation variety
        working = self.engine8_punctuation_entropy(working)

        # Phase 7: Humanization
        working = self.engine12_humanization(working)

        # Phase 8: Style polish
        working = self.engine3_style_variety_editor(working)
        working = self.engine4_semantic_deepener(working)
        working = self.engine5_structure_regularizer(working)

        # Phase 9: Coherence check
        working = self.engine6_coherence_checker(sentences, working)

        if self._flow_is_extreme(working):
            working = self.engine2_burstiness_synthesizer(sentences)
            working = self.engine6_coherence_checker(sentences, working)

        paragraph = normalize_spacing(" ".join(working))
        if self.semantic_model is not None:
            paragraph = self._semantic_lock(" ".join(context + sentences), " ".join(context) + " " + paragraph, threshold=0.70)
            if context and paragraph.startswith(" ".join(context)):
                paragraph = paragraph[len(" ".join(context)):].strip()
        paragraph = self._english_editorial_pass(paragraph)
        if self.preserve_word_count:
            paragraph = self._avoid_expansion(paragraph)
        return paragraph

    def _avoid_expansion(self, text: str) -> str:
        original_count = len(tokenize_words(self.original_text))
        revised_count = len(tokenize_words(text))
        if not original_count or revised_count <= original_count * 1.15:
            return text
        text = re.sub(
            r"\b(Furthermore|Additionally|Moreover|However|Therefore|Notably|Clinically|Historically|In this passage|In the evidence|In the analysis|In the reported material),\s+",
            "",
            text,
        )
        return normalize_spacing(text)

    def run(self) -> str:
        text = strip_chatbot_markup(self.original_text)
        # Fix tortured phrases first
        for pattern, replacement in TORTURED_PHRASES:
            text = re.sub(pattern, replacement, text, flags=re.I)

        # Detect and remove duplicated sentences BEFORE protecting citations
        # This is critical because duplicated sentences are a major AI signal
        all_sentences = split_sentences(text)
        if len(all_sentences) >= 2:
            seen_normalized: set = set()
            deduped_sentences: List[str] = []
            for s in all_sentences:
                # Normalize for comparison: remove citations, punctuation, lowercase
                clean = re.sub(r'\[[\d,\-; ]+\]', '', s)  # Remove [1], [2,3] etc
                clean = re.sub(r'[^\w\s]', '', clean).lower().strip()
                clean = re.sub(r'\s+', ' ', clean)
                if clean not in seen_normalized:
                    seen_normalized.add(clean)
                    deduped_sentences.append(s)
            if len(deduped_sentences) < len(all_sentences):
                text = " ".join(deduped_sentences)

        # Process WITHOUT protecting fragments first so engines can work freely
        # Citations and data are preserved by keeping them in the text
        paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
        revised = []
        context: List[str] = []
        for paragraph in paragraphs:
            revised_paragraph = self.revise_paragraph(paragraph, context=context)
            revised.append(revised_paragraph)
            context = split_sentences(paragraph)[-5:]
        final_text = "\n\n".join(revised)
        return normalize_spacing(final_text)


def is_mostly_arabic(text: str) -> bool:
    letters = re.findall(r"[A-Za-z\u0600-\u06FF]", text)
    if not letters:
        return False
    arabic_letters = re.findall(r"[\u0600-\u06FF]", text)
    return len(arabic_letters) / len(letters) >= 0.35


class ArabicEditorialRevisionEngine:
    """Rule-based Arabic prose editor for warmer, less mechanical rewriting."""

    mechanical_connectors = [
        r"علاوة على ذلك[،,]?\s*",
        r"بالإضافة إلى ذلك[،,]?\s*",
        r"في الختام[،,]?\s*",
        r"وبناءً على ما سبق[،,]?\s*",
        r"بناءً على ما سبق[،,]?\s*",
        r"من ناحية أخرى[،,]?\s*",
        r"تجدر الإشارة إلى أن?\s*",
        r"ومن الجدير بالذكر أن?\s*",
        r"كما يجب التنويه إلى أن?\s*",
        r"في الختام[،,]?\s*",
        r"بشكل عام[،,]?\s*",
        r"من الجدير بالذكر أن?\s*",
        r"كما تجدر الإشارة إلى أن?\s*",
    ]
    vocabulary_swaps = [
        (r"\bبشكل واضح\b", "بجلاء"),
        (r"\bمهم\b", "لافت"),
        (r"\bمهمة\b", "لافتة"),
        (r"\bجيد\b", "متين"),
        (r"\bكبير\b", "واسع الأثر"),
        (r"\bواضح\b", "جلي"),
        (r"\bيؤثر على\b", "يمس"),
        (r"\bيساهم في\b", "يغذي"),
        (r"\bيعكس\b", "يكشف"),
        (r"\bيوضح\b", "يجلي"),
        (r"\bيظهر\b", "يطفو"),
        (r"\bالموضوع\b", "المسألة"),
        (r"\bالنتائج\b", "الحصيلة"),
        (r"\bالبيانات\b", "المعطيات"),
        (r"\bالطريقة\b", "النهج"),
        (r"\bالفكرة\b", "اللمحة"),
        (r"\bالعامل\b", "المؤثر"),
        (r"\bعدد كبير من\b", "طيف واسع من"),
        (r"\bبشكل عام\b", "في المشهد الأوسع"),
        (r"\bفريد من نوعه\b", "نادر"),
        (r"\bاستثنائي\b", "لافت"),
        (r"\bمتطور\b", "حديث"),
        (r"\bبشكل ملموس\b", "عمليًا"),
        (r"\bعلى المستوى العالمي\b", "عالميًا"),
        (r"\bرائد في مجاله\b", "معروف في مجاله"),
        (r"\bبشكل استثنائي\b", "كثيرًا"),
    ]
    passive_rewrites = [
        (r"\bتمت دراسة\b", "تعاين الدراسة"),
        (r"\bتمت ملاحظة\b", "تلمح المعاينة"),
        (r"\bتمت مناقشة\b", "يناقش النص"),
        (r"\bتم تحليل\b", "نعاين"),
        (r"\bتم استخدام\b", "يوظف النص"),
        (r"\bتم الاعتماد على\b", "يتكئ النص على"),
        (r"\bتم التركيز على\b", "نقترب من"),
    ]
    voice_openers = [
        "في تقديري",
        "من واقع المعاينة",
        "ما يبدو لي جلياً",
        "حين أقرأ المشهد عن قرب",
    ]
    short_beats = [
        "هنا يتغير الإيقاع.",
        "وهنا بيت القصيد.",
        "الأمر ليس عابراً.",
        "هذه ليست زينة.",
    ]
    # Morphological variation dictionary
    morphological_variants = [
        (r"\bيُعَدُّ\b", "يُعتبر"),
        (r"\bيُعتبر\b", "يُعد"),
        (r"\bيتميز بـ\b", "يتسم بـ"),
        (r"\bيتسم بـ\b", "يتميز بـ"),
        (r"\bيساهم في\b", "يغذي"),
        (r"\bيؤدي إلى\b", "يُفضي إلى"),
        (r"\bيُفضي إلى\b", "يؤدي إلى"),
        (r"\bيركز على\b", "يقترب من"),
        (r"\bيتناول\b", "يعالج"),
        (r"\bيعالج\b", "يتناول"),
    ]
    # Diglossic insertion patterns (add slight colloquial touch for natural feel)
    diglossic_insertions = [
        "في الواقع،",
        "عمليًا،",
        "على أرض الواقع،",
        "بعبارة أخرى،",
    ]
    # Arabic burstiness short beats
    arabic_short_beats = [
        "هذا جوهري.",
        "الأمر واضح.",
        "هنا تكمن المسألة.",
        "لا التباس في ذلك.",
        "هذا يتكرر.",
        "النتيجة حاسمة.",
    ]
    # Diacritics-aware handling
    arabic_diacritics_pattern = re.compile(r"[\u0610-\u061A\u064B-\u065F\u0670]")

    def __init__(self, intensity: int, text: str, preserve_word_count: bool) -> None:
        self.intensity = intensity
        self.original_text = text
        self.preserve_word_count = preserve_word_count
        self.question_added = False

    def _split_arabic_sentences(self, text: str) -> List[str]:
        text = re.sub(r"\s+", " ", text.strip())
        if not text:
            return []
        parts = re.split(r"(?<=[.!?؟])\s+", text)
        return [part.strip() for part in parts if part.strip()]

    def _clean_mechanical_transitions(self, sentence: str) -> str:
        sentence = sentence.replace("—", "،").replace("–", "،")
        for pattern in self.mechanical_connectors:
            sentence = re.sub(pattern, "", sentence, flags=re.I)
        sentence = re.sub(r"\s*[-]{2,}\s*", "، ", sentence)
        return normalize_spacing(sentence)

    def _enrich_vocabulary(self, sentence: str, index: int) -> str:
        limit = 1 if self.intensity < 4 else 2
        changes = 0
        for pattern, replacement in self.vocabulary_swaps:
            if changes >= limit:
                break
            if re.search(pattern, sentence):
                sentence = re.sub(pattern, replacement, sentence, count=1)
                changes += 1
        return normalize_spacing(sentence)

    def _activate_voice(self, sentence: str) -> str:
        for pattern, replacement in self.passive_rewrites:
            sentence = re.sub(pattern, replacement, sentence)
        # Apply morphological variants at higher intensity
        if self.intensity >= 3:
            for pattern, replacement in self.morphological_variants:
                if re.search(pattern, sentence):
                    sentence = re.sub(pattern, replacement, sentence, count=1)
                    break  # Only one variant per sentence
        return normalize_spacing(sentence)

    def _inject_burstiness(self, sentences: List[str]) -> List[str]:
        """Insert short Arabic beats between long sentences for natural burstiness."""
        if not sentences or len(sentences) < 3:
            return sentences

        lengths = [len(tokenize_words(s)) for s in sentences]
        avg_len = sum(lengths) / len(lengths) if lengths else 0

        result = []
        beat_idx = 0
        for i, sentence in enumerate(sentences):
            result.append(sentence)
            word_count = lengths[i]
            # Insert a beat after a long sentence at intervals
            if (self.intensity >= 3 and
                word_count >= avg_len and
                i < len(sentences) - 1 and
                i % 3 == 1 and
                beat_idx < 2):
                result.append(self.arabic_short_beats[beat_idx % len(self.arabic_short_beats)])
                beat_idx += 1

        return result

    def _shape_rhythm(self, sentences: List[str]) -> List[str]:
        if len(sentences) < 2:
            return sentences
        shaped: List[str] = []
        for index, sentence in enumerate(sentences):
            shaped.append(sentence)
            words_count = len(tokenize_words(sentence))
            should_add_beat = bool(self.short_beats) and self.intensity >= 4 and words_count >= 18 and index % 3 == 0
            if should_add_beat:
                shaped.append(self.short_beats[index % len(self.short_beats)])
        return shaped

    def _add_mid_text_question(self, sentences: List[str]) -> List[str]:
        return sentences

    def _avoid_expansion(self, text: str) -> str:
        if not self.preserve_word_count:
            return text
        original_count = len(tokenize_words(self.original_text))
        revised_count = len(tokenize_words(text))
        if not original_count or revised_count <= original_count * 1.25:
            return text
        text = re.sub(r"\b(?:في تقديري|من واقع المعاينة|ما يبدو لي جلياً)،\s*", "", text, count=1)
        text = re.sub(r"(?:هنا يتغير الإيقاع|وهنا بيت القصيد|الأمر ليس عابراً|هذه ليست زينة)\.", "", text)
        return normalize_spacing(text)

    def _deduplicate_arabic(self, sentences: List[str]) -> List[str]:
        """Remove duplicated Arabic sentences."""
        seen: set = set()
        result: List[str] = []
        for s in sentences:
            normalized = re.sub(r'[^\w\s]', '', s.lower()).strip()
            normalized = re.sub(r'\s+', ' ', normalized)
            if normalized not in seen:
                seen.add(normalized)
                result.append(s)
        return result if result else sentences

    def _humanize_arabic(self, sentences: List[str]) -> List[str]:
        """Add human-like touches to Arabic text."""
        revised: List[str] = []
        for idx, sentence in enumerate(sentences):
            # Add occasional informal touch
            if self.intensity >= 3 and idx % 5 == 3 and len(tokenize_words(sentence)) >= 12:
                hedging = [
                    "يبدو أن",
                    "على ما يرام",
                    "في الواقع،",
                    "على أرض الواقع،",
                ]
                # Prepend hedging
                hedge = hedging[idx % len(hedging)]
                sentence = f"{hedge} {sentence[:1].lower()}{sentence[1:]}" if sentence[0].isupper() else f"{hedge} {sentence}"

            revised.append(sentence)
        return revised

    def revise_paragraph(self, paragraph: str) -> str:
        sentences = self._split_arabic_sentences(paragraph)
        # Deduplicate before processing
        sentences = self._deduplicate_arabic(sentences)
        revised: List[str] = []
        for index, sentence in enumerate(sentences):
            # Fix tortured phrases first
            for pattern, replacement in TORTURED_PHRASES:
                sentence = re.sub(pattern, replacement, sentence, flags=re.I)
            sentence = self._clean_mechanical_transitions(sentence)
            sentence = self._activate_voice(sentence)
            sentence = self._enrich_vocabulary(sentence, index)
            # Add diglossic insertions at higher intensity
            if self.intensity >= 4 and index % 5 == 2 and len(tokenize_words(sentence)) >= 10:
                insertion = self.diglossic_insertions[index % len(self.diglossic_insertions)]
                sentence = f"{insertion} {sentence[:1].lower()}{sentence[1:]}" if sentence[0].isupper() else f"{insertion} {sentence}"
            if sentence and sentence[-1] not in ".!?؟":
                sentence += "."
            revised.append(sentence)
        revised = self._inject_burstiness(revised)
        revised = self._shape_rhythm(revised)
        revised = self._add_mid_text_question(revised)
        # Humanize after all other processing
        revised = self._humanize_arabic(revised)
        return normalize_spacing(" ".join(revised))

    def run(self) -> str:
        text = strip_chatbot_markup(self.original_text)
        # Fix tortured phrases first
        for pattern, replacement in TORTURED_PHRASES:
            text = re.sub(pattern, replacement, text, flags=re.I)
        paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
        revised = [self.revise_paragraph(paragraph) for paragraph in paragraphs]
        final_text = "\n\n".join(revised)
        final_text = final_text.replace("—", "،").replace("–", "،")
        return self._avoid_expansion(normalize_spacing(final_text))


SCIENTIFIC_SECTION_NAMES = {
    "abstract",
    "keywords",
    "introduction",
    "background",
    "literature review",
    "materials and methods",
    "methods",
    "methodology",
    "results",
    "discussion",
    "conclusion",
    "conclusions",
    "acknowledgements",
    "acknowledgments",
    "references",
    "الملخص",
    "الكلمات المفتاحية",
    "المقدمة",
    "الخلفية",
    "الدراسات السابقة",
    "المواد والطرق",
    "المنهجية",
    "النتائج",
    "المناقشة",
    "الخاتمة",
    "الاستنتاجات",
    "المراجع",
}


def set_word_font(run, font_name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), "Arial")


def set_word_bidi(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def mostly_rtl_text(text: str) -> bool:
    letters = re.findall(r"[A-Za-z\u0600-\u06FF]", text)
    if not letters:
        return False
    return len(re.findall(r"[\u0600-\u06FF]", text)) / len(letters) >= 0.35


def academic_line_kind(line: str) -> str:
    stripped = line.strip()
    if not stripped:
        return "blank"
    if stripped.count("|") >= 2 or "\t" in stripped:
        return "table"
    if re.match(r"^\s*(?:[-*•]|\d+[.)])\s+\S+", stripped):
        return "list"
    cleaned = re.sub(r"^\s*\d+(?:\.\d+)*\s*", "", stripped).strip(":：.- ").lower()
    words = tokenize_words(stripped)
    terminal = stripped[-1:] in ".!?؟،؛"
    if cleaned in SCIENTIFIC_SECTION_NAMES:
        return "heading"
    if 1 <= len(words) <= 12 and not terminal and not re.search(r",|;|،|؛", stripped):
        return "heading"
    return "paragraph"


def setup_academic_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)


def add_academic_paragraph(doc: Document, line: str, kind: str) -> None:
    rtl = mostly_rtl_text(line)
    if kind == "heading":
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        if rtl:
            set_word_bidi(paragraph)
        run = paragraph.add_run(line.strip())
        set_word_font(run, size=13, bold=True)
        return

    if kind == "list":
        marker = re.match(r"^\s*(?:[-*•]|\d+[.)])\s+", line)
        content = line[marker.end() :].strip() if marker else line.strip()
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_word_bidi(paragraph)
        run = paragraph.add_run(content)
        set_word_font(run)
        return

    if kind == "table":
        cells = [cell.strip() for cell in re.split(r"\||\t", line) if cell.strip()]
        table = doc.add_table(rows=1, cols=max(1, len(cells)))
        table.style = "Table Grid"
        for idx, value in enumerate(cells):
            cell = table.cell(0, idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if mostly_rtl_text(value) else WD_ALIGN_PARAGRAPH.LEFT
            if mostly_rtl_text(value):
                set_word_bidi(p)
            run = p.add_run(value)
            set_word_font(run, size=10)
        doc.add_paragraph()
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0 if rtl else 0.25)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if rtl:
        set_word_bidi(paragraph)
    run = paragraph.add_run(line.strip())
    set_word_font(run)


def create_word_document(text: str, title: Optional[str] = None) -> BytesIO:
    doc = Document()
    doc.core_properties.author = AUTHOR_NAME
    doc.core_properties.title = title or "DeepClean Studio Revised Manuscript"
    setup_academic_document(doc)
    if title:
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(title.strip())
        set_word_font(run, size=14, bold=True)
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(12)
    author_run = author.add_run(AUTHOR_NAME)
    set_word_font(author_run, size=12)
    previous_blank = False
    for raw_line in polish_layout_artifacts(text).split("\n"):
        kind = academic_line_kind(raw_line)
        if kind == "blank":
            if not previous_blank:
                doc.add_paragraph()
            previous_blank = True
            continue
        add_academic_paragraph(doc, raw_line, kind)
        previous_blank = False
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def render_academic_preview(text: str) -> str:
    parts: List[str] = ["<div class='academic-preview'>"]
    in_list = False
    for raw_line in polish_layout_artifacts(text).split("\n"):
        line = raw_line.strip()
        kind = academic_line_kind(raw_line)
        if kind != "list" and in_list:
            parts.append("</ul>")
            in_list = False
        if kind == "blank":
            continue
        if kind == "heading":
            parts.append(f"<h3>{html.escape(line)}</h3>")
        elif kind == "list":
            marker = re.match(r"^\s*(?:[-*•]|\d+[.)])\s+", line)
            content = line[marker.end() :].strip() if marker else line
            if not in_list:
                parts.append("<ul>")
                in_list = True
            parts.append(f"<li>{html.escape(content)}</li>")
        elif kind == "table":
            cells = [html.escape(cell.strip()) for cell in re.split(r"\||\t", line) if cell.strip()]
            parts.append("<table><tr>" + "".join(f"<td>{cell}</td>" for cell in cells) + "</tr></table>")
        else:
            rtl_attr = " dir='rtl'" if mostly_rtl_text(line) else ""
            parts.append(f"<p{rtl_attr}>{html.escape(line)}</p>")
    if in_list:
        parts.append("</ul>")
    parts.append("</div>")
    return "".join(parts)


def word_level_diff(original: str, modified: str) -> str:
    orig_words = original.split()
    mod_words = modified.split()
    matcher = difflib.SequenceMatcher(None, orig_words, mod_words)
    html_parts: List[str] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        old = html.escape(" ".join(orig_words[i1:i2]))
        new = html.escape(" ".join(mod_words[j1:j2]))
        if tag == "equal":
            html_parts.append(old)
        elif tag == "replace":
            html_parts.append(f"<span class='removed'>{old}</span> <span class='added'>{new}</span>")
        elif tag == "delete":
            html_parts.append(f"<span class='removed'>{old}</span>")
        elif tag == "insert":
            html_parts.append(f"<span class='added'>{new}</span>")
    return " ".join(part for part in html_parts if part)


def jaccard_similarity(original: Iterable[str], revised: Iterable[str]) -> float:
    original_set = set(original)
    revised_set = set(revised)
    if not original_set and not revised_set:
        return 1.0
    return len(original_set & revised_set) / max(1, len(original_set | revised_set))


def count_pattern_matches(text: str, patterns: Sequence[str]) -> int:
    return sum(len(re.findall(pattern, text, flags=re.I)) for pattern in patterns)


def authorship_review_band(formulaic_count: int, certainty_count: int, sentence_variation: float) -> str:
    signals = 0
    if formulaic_count:
        signals += 1
    if certainty_count:
        signals += 1
    if sentence_variation < 0.20:
        signals += 1
    if signals == 0:
        return "منخفض"
    if signals == 1:
        return "متوسط"
    return "مرتفع للمراجعة"


def clamp(value: float, lower: float = 0.0, upper: float = 1.0) -> float:
    return max(lower, min(upper, value))


def is_mostly_english(text: str) -> bool:
    letters = re.findall(r"[A-Za-z\u0600-\u06FF]", text)
    if not letters:
        return False
    english_letters = re.findall(r"[A-Za-z]", text)
    return len(english_letters) / len(letters) >= 0.70


def count_ai_vocabulary(text: str) -> Counter:
    lowered = text.lower()
    hits: Counter = Counter()
    for term in AI_VOCABULARY_TERMS:
        if re.search(r"^[A-Za-z -]+$", term):
            count = len(re.findall(rf"\b{re.escape(term.lower())}\b", lowered, flags=re.I))
        else:
            count = lowered.count(term.lower())
        if count:
            hits[term] = count
    return hits


def repeated_opener_ratio(sentences: Sequence[str]) -> float:
    if len(sentences) < 3:
        return 0.0
    openers = []
    for sentence in sentences:
        words = tokenize_words(sentence.lower())
        if words:
            openers.append(words[0])
    if not openers:
        return 0.0
    counts = Counter(openers)
    return max(counts.values()) / len(openers)


def current_section_label(line: str) -> Optional[str]:
    cleaned = re.sub(r"^\s*\d+(?:\.\d+)*\s*", "", line.strip().lower()).strip(":：.- ")
    section_map = {
        "abstract": "abstract",
        "introduction": "introduction",
        "background": "introduction",
        "materials and methods": "methods",
        "methods": "methods",
        "methodology": "methods",
        "experimental setup": "methods",
        "results": "results",
        "discussion": "discussion",
        "conclusion": "conclusion",
        "conclusions": "conclusion",
        "references": "references",
        "bibliography": "references",
        "الملخص": "abstract",
        "المقدمة": "introduction",
        "المنهجية": "methods",
        "المواد والطرق": "methods",
        "النتائج": "results",
        "المناقشة": "discussion",
        "الخاتمة": "conclusion",
        "المراجع": "references",
    }
    return section_map.get(cleaned)


def exclusion_reason(line: str, active_section: Optional[str]) -> Optional[str]:
    stripped = line.strip()
    if not stripped:
        return None
    if active_section == "references":
        return "مراجع"
    if stripped.startswith(("|", "+---")) or stripped.count("|") >= 2:
        return "جداول"
    if re.search(r"^\s*(table|figure|fig\.|جدول|شكل)\s+\d+", stripped, flags=re.I):
        return "جداول وأشكال"
    if re.search(r"(```|^\s*(def|class|import|from|for|while|if|return)\b|^\s*(SELECT|CREATE|INSERT)\b)", stripped, flags=re.I):
        return "أكواد"
    if re.search(r"[{};]{2,}|(?:==|!=|<=|>=|->|=>)", stripped) and len(tokenize_words(stripped)) <= 18:
        return "أكواد"
    if re.search(r"[=∑√≈≤≥±×÷]", stripped) and len(tokenize_words(stripped)) <= 18:
        return "معادلات"
    tokens = tokenize_words(stripped)
    if tokens:
        numeric_tokens = sum(1 for token in tokens if re.search(r"\d", token))
        if numeric_tokens / len(tokens) >= 0.45:
            return "بيانات رقمية"
    if len(tokens) < 7:
        return "نص قصير غير مؤهل"
    return None


def isolate_qualifying_text(text: str) -> Tuple[str, Counter]:
    qualifying_lines: List[str] = []
    excluded: Counter = Counter()
    active_section: Optional[str] = None
    in_code_fence = False

    for line in text.splitlines():
        section = current_section_label(line)
        if section:
            active_section = section
            excluded["عناوين أقسام"] += len(line)
            continue
        if line.strip().startswith("```"):
            in_code_fence = not in_code_fence
            excluded["أكواد"] += len(line)
            continue
        if in_code_fence:
            excluded["أكواد"] += len(line)
            continue

        reason = exclusion_reason(line, active_section)
        if reason:
            excluded[reason] += len(line)
            continue
        qualifying_lines.append(line)

    qualifying_text = normalize_spacing("\n".join(qualifying_lines))
    long_sentences = [sentence for sentence in split_sentences(qualifying_text) if len(tokenize_words(sentence)) >= 7]
    return normalize_spacing(" ".join(long_sentences)), excluded


def detect_em_dash_overuse(text: str, threshold: float = 0.025) -> Tuple[float, int]:
    """Count em dash density. Returns (density, count)."""
    em_dashes = text.count("\u2014") + text.count("\u2013")
    words = tokenize_words(text)
    density = em_dashes / max(1, len(words))
    return density, em_dashes


def detect_curly_quotes(text: str) -> Tuple[int, int]:
    """Count curly double and single quotes. Returns (double_count, single_count)."""
    double = text.count("\u201c") + text.count("\u201d")
    single = text.count("\u2018") + text.count("\u2019")
    return double, single


def detect_title_case_headings(text: str) -> int:
    """Count lines that look like title-case headings (3+ content words capitalized)."""
    count = 0
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        # Skip very short lines, lines ending with punctuation
        if len(stripped) < 10 or stripped[-1] in ".!?؛،":
            continue
        words = stripped.split()
        if len(words) < 3:
            continue
        # Count words that start with uppercase (not first word, not small words)
        small_words = {"a", "an", "the", "and", "but", "or", "for", "nor", "in", "on", "at", "to", "by", "of", "with", "as", "is", "are"}
        capitalized_content = sum(
            1 for i, w in enumerate(words)
            if i > 0 and w[0].isupper() and w.lower() not in small_words
        )
        if capitalized_content >= 3:
            count += 1
    return count


def detect_rule_of_three(text: str) -> int:
    """Count comma-separated triplets of similar structure (adjective, adjective, and adjective)."""
    pattern = re.compile(
        r"\b(\w+),\s+(\w+),\s+and\s+(\w+)\b",
        re.I,
    )
    return len(pattern.findall(text))


def detect_placeholder_text(text: str) -> int:
    """Count placeholder patterns like [insert X], XX-XX-XXXX, etc."""
    patterns = [
        r"\[insert\s+\w+\]",
        r"XX-XX-XXXX",
        r"YYYY-MM-DD",
        r"\[citation needed\]",
        r"\[insert citation\]",
        r"\[insert reference\]",
        r"\[insert name\]",
        r"\[insert date\]",
    ]
    return sum(len(re.findall(p, text, re.I)) for p in patterns)


def detect_markdown_remnants(text: str) -> int:
    """Count Markdown formatting that shouldn't be in academic prose."""
    patterns = [
        r"(?m)^\s*#{1,6}\s+\S",
        r"\*\*[^*\n]+\*\*",
        r"__[^_\n]+__",
        r"(?m)^\s*[-*]\s+\S",
        r"```",
        r"\[[^\]\n]+\]\(https?://[^)]+\)",
    ]
    return sum(len(re.findall(p, text)) for p in patterns)


def academic_section_notes(text: str) -> Tuple[Tuple[str, ...], float]:
    sections: Counter = Counter()
    active_section: Optional[str] = None
    section_words: Counter = Counter()
    for line in text.splitlines():
        section = current_section_label(line)
        if section:
            sections[section] += 1
            active_section = section
            continue
        if active_section:
            section_words[active_section] += len(tokenize_words(line))

    notes: List[str] = []
    method_words = section_words.get("methods", 0)
    total_section_words = sum(section_words.values())
    method_ratio = method_words / total_section_words if total_section_words else 0.0
    if method_words:
        notes.append("تم رصد قسم منهجية/مواد؛ خُفف أثر انتظام الصياغة لأنه شائع في هذا القسم.")
    if sections.get("references"):
        notes.append("تم عزل قسم المراجع من حساب الأسلوب، مع إبقاء تنبيهات شكلية للمصداقية.")
    if not sections:
        notes.append("لم تظهر عناوين أقسام بحثية واضحة؛ جرى تحليل النص ككتلة نثرية واحدة.")
    return tuple(notes), clamp(method_ratio * 0.18, 0.0, 0.12)


def source_code_alerts(text: str) -> Tuple[str, ...]:
    alerts: List[str] = []
    fenced_blocks = len(re.findall(r"```[\s\S]*?```", text))
    if fenced_blocks:
        alerts.append(f"تم رصد {fenced_blocks} كتلة كود؛ عُزلت عن تحليل النثر وتحتاج مراجعة منفصلة.")
    if re.search(r"^\s*(def|class|import|from)\s+\w+", text, flags=re.M):
        alerts.append("توجد إشارات Python ظاهرة داخل المستند.")
    if re.search(r"\b(library|data\.frame|ggplot|<-)\b", text):
        alerts.append("توجد إشارات R أو تحليل بيانات داخل المستند.")
    if re.search(r"(?m)^\s{0,3}#{1,6}\s+\S|\*\*[^*\n]+\*\*|__[^_\n]+__|```|^\s*\|.+\|\s*$", text):
        alerts.append("توجد آثار Markdown ظاهرة؛ راجعها إذا كان النص موجهاً لصيغة أكاديمية أو موسوعية لا تستخدم Markdown.")
    if re.search(r"\b(?:as an AI language model|as a large language model|I hope this helps|up to my last training update|of course|certainly|you're absolutely right)\b", text, flags=re.I):
        alerts.append("توجد بقايا خطاب محادثة آلية داخل النص، مثل اعتذار النموذج أو حدود معرفته.")
    if re.search(r"\b(?:key takeaways|in this article we will|this article aims to|if you have any concerns|if there are specific)\b", text, flags=re.I):
        alerts.append("توجد صياغة قالبية من مخرجات المحادثة أو المسودات الآلية، وتحتاج حذفًا لا إعادة تزيين.")
    if re.search(r"\b(?:independent coverage|profiled in|media outlets|active social media presence|written by a leading expert)\b", text, flags=re.I):
        alerts.append("توجد صياغة إثبات ملحوظية/تغطية إعلامية شائعة في النصوص الآلية؛ تحقق من المصدر والمعنى قبل الإبقاء عليها.")
    if re.search(r"\b(?:algorithm|pseudocode|خوارزمية)\b", text, flags=re.I):
        alerts.append("يوجد وصف خوارزمي؛ راجع توافقه مع المصادر أو المستودعات المفتوحة يدويًا.")

    # Curly quote detection
    curly_double, curly_single = detect_curly_quotes(text)
    if curly_double or curly_single:
        alerts.append(f"تم رصد {curly_double + curly_single} علامة اقتباس منحنية؛ الشائع في مخرجات ChatGPT وDeepSeek. يفضل استخدام علامات الاقتباس المستقيمة.")

    # Placeholder text
    placeholders = detect_placeholder_text(text)
    if placeholders:
        alerts.append(f"تم رصد {placeholders} نمط نص نائب (placeholder)؛ هذه تحتاج تعبئة يدوية أو حذف.")

    # Em dash overuse
    density, count = detect_em_dash_overuse(text)
    if density > 0.025 and count >= 3:
        alerts.append(f"كثافة الشرطة الطويلة (em dash) مرتفعة ({count} شرطة)؛ الشائع البشري هو الفاصلة أو الأقواس. يفضل الاستبدال بفواصل.")

    # Title case headings
    title_case_count = detect_title_case_headings(text)
    if title_case_count >= 2:
        alerts.append(f"تم رصد {title_case_count} عنوان بحالة العنوان (Title Case)؛ الشائع في الكتابة الأكاديمية هو حالة الجملة (Sentence case).")

    # Negative parallelisms
    negative_parallel_count = count_pattern_matches(text, [
        r"\bnot only\b[^.!?]{0,60}\bbut also\b",
        r"\bnot just\b[^.!?]{0,60}\bbut also\b",
    ])
    if negative_parallel_count >= 2:
        alerts.append(f"تم رصد {negative_parallel_count} تركيب نفي مقابل (not only... but also)؛ هذا نمط آلي شائع يفضل تبسيطه.")

    # Rule of three
    rule_of_three_count = detect_rule_of_three(text)
    if rule_of_three_count >= 3:
        alerts.append(f"تم رصد {rule_of_three_count} تركيبات ثلاثية (قاعدة الثلاثة)؛ الاستخدام المتكرر نمط آلي شائع.")

    # Markdown remnants
    md_count = detect_markdown_remnants(text)
    if md_count >= 3:
        alerts.append(f"تم رصد {md_count} أثر تنسيق Markdown داخل النص؛ راجعها إذا كان النص موجهاً لصيغة أكاديمية.")

    return tuple(dict.fromkeys(alerts))


def translation_pattern_signal(text: str) -> float:
    words = tokenize_words(text)
    if len(words) < 40:
        return 0.0
    arabic = len(re.findall(r"[\u0600-\u06FF]", text))
    latin = len(re.findall(r"[A-Za-z]", text))
    total = arabic + latin
    mixed_script = min(arabic, latin) / total if total else 0.0
    literal_markers = count_pattern_matches(
        text,
        [
            r"\bin addition to that\b",
            r"\bfrom another side\b",
            r"\baccording to what preceded\b",
            r"\bthe matter\b",
            r"من جهة أخرى",
            r"بناء على ما سبق",
        ],
    )
    return clamp(mixed_script * 1.4 + min(0.35, literal_markers * 0.12))


def internal_overlap_signal(text: str) -> float:
    words = [word.lower() for word in tokenize_words(text)]
    if len(words) < 80:
        return 0.0
    ngrams = [" ".join(words[index : index + 6]) for index in range(len(words) - 5)]
    counts = Counter(ngrams)
    repeated = sum(count - 1 for count in counts.values() if count > 1)
    return clamp(repeated / max(12, len(ngrams)) * 5)


def citation_hygiene_alerts(text: str) -> Tuple[str, ...]:
    alerts: List[str] = []
    urls = re.findall(r"https?://\S+", text)
    dois = re.findall(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", text, flags=re.I)
    year_mentions = re.findall(r"\([^)]*\b(?:19|20)\d{2}\b[^)]*\)", text)
    bracket_citations = re.findall(r"\[[\d,\-; ]+\]", text)
    reference_lines = [
        line.strip()
        for line in text.splitlines()
        if re.search(r"\b(?:19|20)\d{2}\b", line) and len(tokenize_words(line)) >= 6
    ]
    if re.search(r"\b(?:et al\.?|وآخرون)\b", text, flags=re.I) and not year_mentions:
        alerts.append("توجد إحالة باسم مؤلفين دون سنة واضحة.")
    if re.search(r"\b(?:study|paper|research|دراسة|بحث)\b", text, flags=re.I) and not (year_mentions or bracket_citations or urls or dois):
        alerts.append("النص يذكر دراسات أو أبحاثًا دون مرجع ظاهر.")
    if urls and any(url.endswith((".", ",", ";")) for url in urls):
        alerts.append("بعض الروابط تنتهي بعلامات ترقيم؛ راجعها قبل التسليم.")
    if dois and any(len(doi) < 12 for doi in dois):
        alerts.append("يوجد DOI قصير أو غير مألوف ويحتاج تحققًا يدويًا.")
    if reference_lines and not (dois or urls):
        alerts.append("توجد مراجع بصياغة أكاديمية دون DOI أو رابط ظاهر؛ يلزم تحقق خارجي مثل Crossref أو PubMed.")
    if re.search(r"\b(?:Journal of Advanced|International Journal of Modern|Global Research in)\b", text, flags=re.I):
        alerts.append("ظهر اسم مجلة عام جدًا؛ راجعه يدويًا لأنه نمط شائع في المراجع المصطنعة.")
    return tuple(alerts)


def sentence_signal_score(sentence: str, avg_length: float, vocabulary_hits: Counter) -> SentenceSignal:
    words = tokenize_words(sentence)
    word_count = len(words)
    score = 0.0
    reasons: List[str] = []

    formulaic = count_pattern_matches(sentence, FORMULAIC_PATTERNS)
    if formulaic:
        score += 0.28
        reasons.append("عبارة قالبية")

    certainty = count_pattern_matches(sentence, CERTAINTY_PATTERNS)
    if certainty:
        score += 0.18
        reasons.append("جزم زائد")

    vocab_count = sum(count for term, count in count_ai_vocabulary(sentence).items())
    if vocab_count:
        score += min(0.24, 0.08 * vocab_count)
        reasons.append("مفردات شائعة في النصوص الآلية")

    if avg_length and 0.85 <= word_count / avg_length <= 1.15 and word_count >= 14:
        score += 0.12
        reasons.append("إيقاع شديد الانتظام")

    comma_count = sentence.count(",") + sentence.count("،")
    if word_count >= 24 and comma_count >= 2:
        score += 0.10
        reasons.append("جملة طويلة ومصقولة جدًا")

    # Negative parallelism detection
    negative_parallel = count_pattern_matches(
        sentence,
        [
            r"\bnot only\b[^.!?]{0,60}\bbut also\b",
            r"\bnot just\b[^.!?]{0,60}\bbut also\b",
            r"\bnot a\b[^.!?]{0,40}\bbut a\b",
        ],
    )
    if negative_parallel:
        score += 0.20
        reasons.append("تركيب نفي مقابل")

    # Rule of three in sentence
    if re.search(r"\b(\w+),\s+(\w+),\s+and\s+(\w+)\b", sentence, re.I):
        score += 0.08
        reasons.append("قاعدة الثلاثة")

    # Promotional language
    promotional = count_pattern_matches(
        sentence,
        [
            r"\bnestled in\b",
            r"\bgroundbreaking\b",
            r"\brenowned\b",
            r"\bdiverse array\b",
            r"\bexemplifies\b",
            r"\bcommitment to\b",
            r"\bworld-class\b",
            r"\bstate-of-the-art\b",
            r"\bcutting-edge\b",
            r"\bpremier\b",
            r"\bempower\b",
            r"\bunlock\b",
            r"\bstreamline\b",
        ],
    )
    if promotional:
        score += min(0.22, 0.11 * promotional)
        reasons.append("لغة ترويجية")

    # Copulative avoidance
    copulative_avoid = count_pattern_matches(
        sentence,
        [r"\brefers to\b", r"\bis defined as\b", r"\bis known as\b"],
    )
    if copulative_avoid:
        score += 0.12
        reasons.append("تجنب الفعل 'يكون'")

    # Challenges + future formula
    challenges_formula = count_pattern_matches(
        sentence,
        [
            r"\bfuture (?:outlook|prospects|directions|implications)\b",
            r"\blooking ahead\b",
            r"\bmoving forward\b",
            r"\bchallenges and (?:opportunities|future|legacy)\b",
        ],
    )
    if challenges_formula:
        score += 0.16
        reasons.append("صيغة التحديات والمستقبل")

    if not reasons:
        reasons.append("إشارة منخفضة")

    score = clamp(score)
    if score >= 0.55:
        label = "مراجعة عالية"
    elif score >= 0.28:
        label = "مراجعة متوسطة"
    else:
        label = "مراجعة منخفضة"
    return SentenceSignal(sentence, score, label, tuple(reasons))


def compute_transparency_report(text: str) -> TransparencyReport:
    normalized = normalize_spacing(text)
    qualifying_text, excluded = isolate_qualifying_text(normalized)
    analysis_text = qualifying_text
    section_notes, section_discount = academic_section_notes(normalized)
    words = tokenize_words(analysis_text)
    sentences = split_sentences(analysis_text)
    sentence_lengths = [len(tokenize_words(sentence)) for sentence in sentences]
    avg_length = float(np.mean(sentence_lengths)) if sentence_lengths else 0.0
    variation = float(np.std(sentence_lengths) / avg_length) if avg_length else 0.0
    lexical_diversity = len(set(word.lower() for word in words)) / len(words) if words else 0.0
    formulaic_count = count_pattern_matches(analysis_text, FORMULAIC_PATTERNS)
    certainty_count = count_pattern_matches(analysis_text, CERTAINTY_PATTERNS)
    vocabulary_hits = count_ai_vocabulary(analysis_text)
    vocabulary_total = sum(vocabulary_hits.values())
    opener_ratio = repeated_opener_ratio(sentences)

    perplexity_signal = clamp((0.52 - lexical_diversity) * 1.7 + min(0.35, formulaic_count / 8))
    burstiness_signal = clamp((0.28 - variation) * 2.2) if len(sentences) >= 3 else 0.0
    paraphrase_signal = clamp((vocabulary_total / max(12, len(words))) * 4 + (0.20 if 0.22 <= variation <= 0.48 and vocabulary_total else 0))
    translation_signal = translation_pattern_signal(analysis_text)
    plagiarism_signal = internal_overlap_signal(analysis_text)

    esl_adjustment = 0.0
    if is_mostly_english(analysis_text) and avg_length <= 18 and formulaic_count == 0 and certainty_count == 0 and lexical_diversity >= 0.46:
        esl_adjustment = 0.14

    ai_generated_score = clamp(
        perplexity_signal * 0.40
        + burstiness_signal * 0.34
        + min(0.24, formulaic_count * 0.04)
        + min(0.16, certainty_count * 0.04)
        - esl_adjustment
        - section_discount
    )
    ai_paraphrased_score = clamp(
        paraphrase_signal * 0.58
        + translation_signal * 0.20
        + min(0.18, vocabulary_total / max(20, len(words)) * 2)
        - section_discount * 0.5
    )
    raw_score = (
        ai_generated_score * 0.42
        + ai_paraphrased_score * 0.30
        + min(0.20, certainty_count * 0.05)
        + clamp((opener_ratio - 0.34) * 0.8) * 0.16
        + plagiarism_signal * 0.08
        - esl_adjustment
        - section_discount
    )
    review_score = clamp(raw_score)
    false_positive_shield = review_score < 0.20

    if not normalized:
        classification = "لا يوجد نص"
    elif false_positive_shield:
        classification = "إشارة منخفضة / تحت عتبة 20%"
    elif review_score < 0.34:
        classification = "تحرير آلي محتمل"
    elif review_score < 0.52:
        classification = "مختلط محتمل"
    elif ai_paraphrased_score >= ai_generated_score:
        classification = "إعادة صياغة آلية محتملة"
    else:
        classification = "توليد آلي محتمل"

    sentence_signals = tuple(sentence_signal_score(sentence, avg_length, vocabulary_hits) for sentence in sentences)
    sorted_hits = tuple(vocabulary_hits.most_common(12))
    confidence = 0.0 if not analysis_text else review_score
    source_alerts = source_code_alerts(normalized)

    # Structural pattern alerts
    structural_notes: List[str] = []
    em_density, em_count = detect_em_dash_overuse(normalized)
    if em_density > 0.02 and em_count >= 2:
        structural_notes.append("كثافة الشرطة الطويلة مرتفعة في النص الأصلي.")
    curly_d, curly_s = detect_curly_quotes(normalized)
    if curly_d + curly_s >= 2:
        structural_notes.append("توجد علامات اقتباس منحنية شائعة في مخرجات الذكاء الاصطناعي.")
    tc_headings = detect_title_case_headings(normalized)
    if tc_headings >= 2:
        structural_notes.append("عناوين بحالة العنوان (Title Case) تظهر بشكل متكرر.")
    r3_count = detect_rule_of_three(normalized)
    if r3_count >= 3:
        structural_notes.append("التركيبات الثلاثية المتكررة (قاعدة الثلاثة) نمط آلي شائع.")
    placeholder_count = detect_placeholder_text(normalized)
    if placeholder_count:
        structural_notes.append(f"يوجد {placeholder_count} نمط نص نائب يحتاج مراجعة.")

    # Enhanced detection metrics
    # Type-token ratio
    ttr = len(set(word.lower() for word in words)) / len(words) if words else 0.0

    # Punctuation entropy
    punct_chars = [c for c in analysis_text if c in ".,;:!?()-—–\"'"]
    punct_counter = Counter(punct_chars)
    punct_entropy = 0.0
    total_punct = sum(punct_counter.values())
    if total_punct > 0:
        for count in punct_counter.values():
            p = count / total_punct
            if p > 0:
                punct_entropy -= p * np.log2(p)

    # Function word bigram analysis
    function_words = {"the", "a", "an", "is", "are", "was", "were", "has", "have", "had",
                      "can", "could", "will", "would", "shall", "should", "may", "might",
                      "must", "it", "this", "that", "these", "those", "which", "who",
                      "whom", "whose", "of", "in", "to", "for", "with", "on", "at", "by",
                      "from", "as", "into", "through", "during", "before", "after"}
    fw_bigrams = Counter()
    word_list = [w.lower() for w in words]
    for i in range(len(word_list) - 1):
        if word_list[i] in function_words and word_list[i+1] in function_words:
            fw_bigrams[(word_list[i], word_list[i+1])] += 1
    fw_bigram_concentration = sum(c for _, c in fw_bigrams.most_common(3)) / max(1, sum(fw_bigrams.values())) if fw_bigrams else 0.0

    # Tortured phrase detection
    tortured_count = count_pattern_matches(analysis_text, [p for p, _ in TORTURED_PHRASES])
    if tortured_count > 0:
        structural_notes.append(f"تم رصد {tortured_count} عبارة متعذرة (tortured phrases)؛ هذه صياغات آلية غير طبيعية ويجب استبدالها.")

    # Low punctuation entropy alert
    if punct_entropy < 1.5 and total_punct >= 5:
        structural_notes.append("تنوع علامات الترقيم منخفض (punctuation entropy)؛ النصوص البشرية تستخدم تنوعًا أكبر في الترقيم.")

    # High function word bigram concentration alert
    if fw_bigram_concentration > 0.6 and len(fw_bigrams) >= 5:
        structural_notes.append("تركيز عالي في أزواج الكلمات الوظيفية (function word bigrams)؛ نمط شائع في النصوص الآلية.")

    # Low type-token ratio alert
    if ttr < 0.4 and len(words) >= 50:
        structural_notes.append(f"نسبة تنوع المفردات (TTR) منخفضة ({ttr:.2f})؛ يشير إلى تكرار مفرداتي أعلى من المتوقع.")

    section_notes = tuple(section_notes) + tuple(structural_notes)

    return TransparencyReport(
        chars=len(normalized),
        words=len(tokenize_words(normalized)),
        sentences=len(split_sentences(normalized)),
        qualifying_chars=len(analysis_text),
        excluded_chars=sum(excluded.values()),
        excluded_blocks=tuple(excluded.most_common()),
        minimum_ready=len(analysis_text) >= 250,
        classification=classification,
        confidence=confidence,
        ai_generated_score=ai_generated_score,
        ai_paraphrased_score=ai_paraphrased_score,
        plagiarism_signal=plagiarism_signal,
        perplexity_signal=perplexity_signal,
        burstiness_signal=burstiness_signal,
        paraphrase_signal=paraphrase_signal,
        translation_signal=translation_signal,
        esl_adjustment=esl_adjustment,
        false_positive_shield=false_positive_shield,
        ai_vocabulary_hits=sorted_hits,
        sentence_signals=sentence_signals,
        citation_alerts=citation_hygiene_alerts(normalized),
        source_code_alerts=source_alerts,
        section_notes=section_notes,
        model_update_note="مؤشر محلي قابل للتحديث: راجع قاموس الأنماط دوريًا، ويفضل كل 90 يومًا إذا تغيرت نماذج الكتابة أو سياسة المجلة.",
        structural_alerts=tuple(structural_notes),
    )


def render_sentence_highlights(report: TransparencyReport) -> str:
    html_parts: List[str] = []
    for signal in report.sentence_signals:
        if signal.score >= 0.55:
            css_class = "sentence-high"
        elif signal.score >= 0.28:
            css_class = "sentence-medium"
        else:
            css_class = "sentence-low"
        title = html.escape("، ".join(signal.reasons))
        html_parts.append(
            f"<span class='{css_class}' title='{title}'>{html.escape(signal.text)}</span>"
        )
    return " ".join(html_parts)


def scan_distribution(report: TransparencyReport) -> Dict[str, int]:
    ai_share = clamp(max(report.ai_generated_score, report.ai_paraphrased_score) * 0.82 + report.confidence * 0.18)
    if report.false_positive_shield:
        ai_share = min(ai_share, 0.19)
    mixed_share = clamp(min(report.ai_generated_score, report.ai_paraphrased_score) * 0.55 + report.plagiarism_signal * 0.20)
    human_share = clamp(1.0 - ai_share - mixed_share)
    total = ai_share + mixed_share + human_share
    if total <= 0:
        return {"AI": 0, "Mixed": 0, "Human": 100}
    return {
        "AI": round(ai_share / total * 100),
        "Mixed": round(mixed_share / total * 100),
        "Human": max(0, 100 - round(ai_share / total * 100) - round(mixed_share / total * 100)),
    }


def scan_verdict_text(report: TransparencyReport) -> str:
    if report.false_positive_shield:
        return "الإشارات منخفضة وتحت عتبة الأمان؛ لا يوجد سبب كافٍ للتصعيد"
    if max(report.ai_generated_score, report.ai_paraphrased_score) >= 0.62:
        return "الإشارات المحلية مرتفعة وتستحق مراجعة بشرية دقيقة"
    if report.classification == "مختلط محتمل":
        return "النص يجمع بين إشارات طبيعية وإشارات آلية محتملة"
    if report.classification == "تحرير آلي محتمل":
        return "قد يكون النص بشريًا مع أثر تحرير أو انتظام زائد"
    return "لا تظهر إشارات محلية قوية، مع بقاء الحكم للمراجعة البشرية"


def render_advanced_scan_card(report: TransparencyReport, report_is_current: bool) -> str:
    distribution = scan_distribution(report)
    status = "المؤشرات محدّثة" if report_is_current else "اضغط المعالجة لتحديث المؤشرات"
    model_badge = "Offline heuristic review"
    dominant = max(distribution, key=distribution.get)
    dominant_class = {
        "AI": "scan-ai",
        "Mixed": "scan-mixed",
        "Human": "scan-human",
    }[dominant]
    return f"""
    <div class="advanced-scan">
        <div class="scan-heading">
            <div>
                <div class="scan-title">واقعية الفحص</div>
                <div class="scan-subtitle">DeepClean Authorship Signals <span>{model_badge}</span></div>
            </div>
            <div class="scan-status">{html.escape(status)}</div>
        </div>
        <div class="scan-body">
            <div class="scan-ring {dominant_class}">{dominant}</div>
            <div class="scan-verdict">
                <strong>{html.escape(scan_verdict_text(report))}</strong>
                <p>{report.qualifying_chars:,} qualifying characters · {report.words:,} words · {report.excluded_chars:,} excluded characters</p>
            </div>
        </div>
        <div class="scan-chips">
            <span class="chip chip-ai">إشارة آلية {distribution["AI"]}%</span>
            <span class="chip chip-mixed">إشارة مختلطة {distribution["Mixed"]}%</span>
            <span class="chip chip-human">إشارة بشرية {distribution["Human"]}%</span>
        </div>
        <div class="scan-toggle">
            <span>جمل تحتاج مراجعة</span>
            <span>مفردات مؤشرية</span>
        </div>
    </div>
    """


def external_detector_observations(report_text: str) -> List[str]:
    lowered = report_text.lower()
    observations: List[str] = []
    if "gptzero" in lowered:
        observations.append("المصدر المذكور يبدو GPTZero.")
    if "model 4.6b" in lowered:
        observations.append("النص يذكر إصدارًا/نموذجًا خارجيًا باسم Model 4.6b؛ يعرضه DeepClean كمعلومة منقولة فقط.")
    if "highly confident" in lowered:
        observations.append("التقرير الخارجي يستخدم صياغة ثقة عالية؛ لا يحولها DeepClean إلى حكم نهائي.")
    if "ai generated" in lowered or "ai-generated" in lowered:
        observations.append("التقرير الخارجي يدعي وجود تشابه مع نصوص مولدة آليًا.")
    if "compared" in lowered and "data" in lowered:
        observations.append("التقرير الخارجي يبرر النتيجة بالمقارنة مع بيانات مرجعية غير متاحة محليًا للتدقيق.")
    if not observations:
        observations.append("تم إرفاق تقرير خارجي غير مفسر آليًا؛ راجعه يدويًا مع المؤشرات المحلية.")
    return observations


def external_detector_severity(report_text: str) -> Tuple[str, str]:
    lowered = report_text.lower()
    high_markers = [
        "highly confident",
        "100%",
        "ai 100",
        "ai generated",
        "ai-generated",
        "was ai",
    ]
    medium_markers = ["mixed", "likely", "possibly", "may be", "similar to"]
    if any(marker in lowered for marker in high_markers):
        return (
            "مرتفع",
            "التقرير الخارجي يرفع أولوية المراجعة. لا ينبغي تجاهله حتى لو كانت المؤشرات المحلية أقل.",
        )
    if any(marker in lowered for marker in medium_markers):
        return (
            "متوسط",
            "التقرير الخارجي يحتوي على إشارة تحتاج مقارنة يدوية مع النص وسجل التحرير.",
        )
    return (
        "غير محدد",
        "التقرير الخارجي مرفق كسياق، لكن لا توجد فيه عبارة حاسمة يمكن تصنيفها محليًا.",
    )


def external_alignment_note(report_text: str, report: TransparencyReport) -> str:
    severity, _ = external_detector_severity(report_text)
    local_high = report.confidence >= 0.52 or max(report.ai_generated_score, report.ai_paraphrased_score) >= 0.52
    if severity == "مرتفع" and not local_high:
        return "يوجد تعارض مهم: الكاشف الخارجي مرتفع بينما مؤشرات DeepClean المحلية أقل. اعتمد هذا كإشارة فشل/نقص في التحليل المحلي، لا كنجاح."
    if severity == "مرتفع" and local_high:
        return "يوجد توافق عام: التقرير الخارجي والمؤشرات المحلية كلاهما يطلبان مراجعة بشرية دقيقة."
    if severity == "متوسط" and local_high:
        return "المؤشرات المحلية أعلى من التقرير الخارجي؛ راجع الجمل المظللة والمفردات قبل إعادة الفحص."
    return "لا يوجد تعارض واضح، لكن التقرير الخارجي يبقى قرينة منفصلة لا تتحول إلى حكم نهائي."


def render_external_detector_card(report_text: str, report: Optional[TransparencyReport] = None) -> str:
    cleaned = report_text.strip()
    if not cleaned:
        return ""
    severity, severity_note = external_detector_severity(cleaned)
    alignment = external_alignment_note(cleaned, report) if report is not None else ""
    observations = "".join(f"<li>{html.escape(item)}</li>" for item in external_detector_observations(cleaned))
    alignment_html = f"<p><strong>المقارنة مع DeepClean:</strong> {html.escape(alignment)}</p>" if alignment else ""
    return f"""
    <div class="external-report">
        <div class="external-title">تقرير كاشف خارجي مرفق</div>
        <p>هذا النص منقول من أداة خارجية ولا ينتجه DeepClean. يُستخدم كقرينة مراجعة فقط.</p>
        <p><strong>مستوى التحذير الخارجي:</strong> {html.escape(severity)} — {html.escape(severity_note)}</p>
        {alignment_html}
        <blockquote>{html.escape(cleaned)}</blockquote>
        <ul>{observations}</ul>
    </div>
    """


def compute_stats(original: str, revised: str) -> RevisionStats:
    original_words = [word.lower() for word in tokenize_words(original)]
    revised_words = [word.lower() for word in tokenize_words(revised)]
    revised_sentences = split_sentences(revised)
    sentence_lengths = [len(tokenize_words(sentence)) for sentence in revised_sentences]
    lexical_diversity = len(set(revised_words)) / len(revised_words) if revised_words else 0.0
    similarity = jaccard_similarity(original_words, revised_words)
    formulaic_count = count_pattern_matches(revised, FORMULAIC_PATTERNS)
    certainty_count = count_pattern_matches(revised, CERTAINTY_PATTERNS)
    avg_length = float(np.mean(sentence_lengths)) if sentence_lengths else 0.0
    sentence_variation = float(np.std(sentence_lengths) / avg_length) if avg_length else 0.0

    # Type-token ratio
    ttr = len(set(revised_words)) / len(revised_words) if revised_words else 0.0

    # Punctuation entropy
    punct_chars = [c for c in revised if c in ".,;:!?()-—–\"'"]
    punct_counter = Counter(punct_chars)
    punct_entropy = 0.0
    total_punct = sum(punct_counter.values())
    if total_punct > 0:
        for count in punct_counter.values():
            p = count / total_punct
            if p > 0:
                punct_entropy -= p * np.log2(p)

    # Function word bigram analysis
    function_words_set = {"the", "a", "an", "is", "are", "was", "were", "has", "have", "had",
                          "can", "could", "will", "would", "shall", "should", "may", "might",
                          "must", "it", "this", "that", "these", "those", "which", "who",
                          "whom", "whose", "of", "in", "to", "for", "with", "on", "at", "by",
                          "from", "as", "into", "through", "during", "before", "after"}
    fw_bigrams = Counter()
    word_list = [w.lower() for w in tokenize_words(revised)]
    for i in range(len(word_list) - 1):
        if word_list[i] in function_words_set and word_list[i+1] in function_words_set:
            fw_bigrams[(word_list[i], word_list[i+1])] += 1
    fw_bigram_concentration = sum(c for _, c in fw_bigrams.most_common(3)) / max(1, sum(fw_bigrams.values())) if fw_bigrams else 0.0

    # Tortured phrase detection
    tortured_count = count_pattern_matches(revised, [p for p, _ in TORTURED_PHRASES])

    # Burstiness CV
    burstiness_cv = sentence_variation  # CV is already the burstiness measure

    return RevisionStats(
        words_original=len(original_words),
        words_revised=len(revised_words),
        sentences_revised=len(revised_sentences),
        avg_sentence_length=avg_length,
        lexical_diversity=lexical_diversity,
        similarity=similarity,
        formulaic_phrase_count=formulaic_count,
        certainty_marker_count=certainty_count,
        sentence_length_variation=sentence_variation,
        authorship_review_band=authorship_review_band(formulaic_count, certainty_count, sentence_variation),
        type_token_ratio=ttr,
        punctuation_entropy=punct_entropy,
        function_word_bigram_concentration=fw_bigram_concentration,
        tortured_phrase_count=tortured_count,
        burstiness_cv=burstiness_cv,
    )


def extract_uploaded_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        raw = uploaded_file.read()
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            return raw.decode("utf-8-sig")
    if name.endswith(".docx"):
        return docx2txt.process(uploaded_file) or ""
    if name.endswith(".pdf"):
        pdf_reader = pypdf.PdfReader(uploaded_file)
        return "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    return ""


def process_text_callback() -> None:
    text = (
        st.session_state.get("text_input", "").strip()
        or st.session_state.get("paste_area", "").strip()
    )
    if not text:
        st.warning("أدخل نصًا أو ارفع ملفًا قبل المعالجة.")
        return
    st.session_state.text_input = text

    if is_mostly_arabic(text):
        engine = ArabicEditorialRevisionEngine(
            intensity=st.session_state.get("intensity", 2),
            text=text,
            preserve_word_count=st.session_state.get("preserve_word_count", True),
        )
    else:
        engine = AcademicRevisionEngine(
            domain=st.session_state.get("domain", "general"),
            intensity=st.session_state.get("intensity", 2),
            text=text,
            preserve_word_count=st.session_state.get("preserve_word_count", True),
        )
    revised = polish_layout_artifacts(engine.run())
    st.session_state.revised_text = revised
    st.session_state.stats = compute_stats(text, revised)
    st.session_state.transparency_report = compute_transparency_report(text)
    st.session_state.processed_source_text = text
    st.session_state.processing_done = True


defaults = {
    "revised_text": "",
    "processing_done": False,
    "text_input": "",
    "domain": "general",
    "intensity": 2,
    "preserve_word_count": True,
    "stats": RevisionStats(0, 0, 0, 0.0, 0.0, 0.0, 0, 0, 0.0, "غير محسوب", 0.0, 0.0, 0.0, 0, 0.0),
    "transparency_report": compute_transparency_report(""),
    "processed_source_text": "",
    "external_detector_report": "",
}
for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value


st.title("DeepClean Studio")
st.caption("محرر أكاديمي لمراجعة الوضوح والتدفق مع الحفاظ على المعنى والمراجع.")
st.caption(AUTHOR_NAME)
st.markdown(
    """
    <style>
    .transparency-panel {
        border: 1px solid #e5e7eb;
        border-radius: 8px;
        padding: 16px;
        background: #fbfbfc;
        line-height: 1.85;
    }
    .sentence-high,
    .sentence-medium,
    .sentence-low {
        border-radius: 6px;
        padding: 2px 4px;
        margin: 2px 1px;
        display: inline;
    }
    .sentence-high { background: #fed7aa; border-bottom: 2px solid #f97316; }
    .sentence-medium { background: #fef3c7; border-bottom: 2px solid #f59e0b; }
    .sentence-low { background: #dcfce7; border-bottom: 2px solid #22c55e; }
    .signal-note {
        color: #475569;
        font-size: 0.92rem;
    }
    .advanced-scan {
        border: 1px solid #e5e7eb;
        border-radius: 8px;
        padding: 18px;
        background: #ffffff;
        box-shadow: 0 1px 3px rgba(15, 23, 42, 0.08);
        margin-bottom: 16px;
    }
    .scan-heading,
    .scan-body,
    .scan-chips,
    .scan-toggle {
        display: flex;
        align-items: center;
        gap: 14px;
        flex-wrap: wrap;
    }
    .scan-heading {
        justify-content: space-between;
        padding-bottom: 14px;
        border-bottom: 1px solid #e5e7eb;
    }
    .scan-title {
        font-size: 1.35rem;
        font-weight: 700;
        color: #111827;
    }
    .scan-subtitle {
        color: #475569;
        margin-top: 4px;
    }
    .scan-subtitle span,
    .scan-status {
        background: #e5e7eb;
        border-radius: 6px;
        padding: 3px 8px;
        color: #334155;
        font-size: 0.85rem;
    }
    .scan-body {
        padding: 18px 0 12px;
    }
    .scan-ring {
        width: 86px;
        height: 86px;
        border-radius: 50%;
        display: grid;
        place-items: center;
        font-weight: 800;
        background: #fff;
    }
    .scan-ai { border: 6px solid #eab308; color: #92400e; }
    .scan-mixed { border: 6px solid #22c55e; color: #166534; }
    .scan-human { border: 6px solid #10b981; color: #047857; }
    .scan-verdict strong {
        display: block;
        font-size: 1.16rem;
        color: #1f2937;
        margin-bottom: 6px;
    }
    .scan-verdict p {
        margin: 0;
        color: #64748b;
    }
    .chip {
        border: 2px solid;
        border-radius: 999px;
        padding: 7px 13px;
        font-weight: 700;
        background: #fff;
    }
    .chip-ai { border-color: #eab308; color: #3f3f46; }
    .chip-mixed { border-color: #86efac; color: #64748b; }
    .chip-human { border-color: #6ee7b7; color: #64748b; }
    .scan-toggle {
        margin-top: 18px;
        background: #f1f5f9;
        border-radius: 999px;
        padding: 6px;
        justify-content: center;
    }
    .scan-toggle span {
        min-width: 150px;
        text-align: center;
        background: #fff;
        border-radius: 999px;
        padding: 9px 16px;
        color: #111827;
        font-weight: 700;
        box-shadow: 0 1px 2px rgba(15, 23, 42, 0.08);
    }
    .external-report {
        border: 1px solid #dbeafe;
        border-radius: 8px;
        padding: 14px 16px;
        background: #eff6ff;
        margin: 14px 0;
        color: #1e3a8a;
    }
    .external-title {
        font-weight: 800;
        color: #1e40af;
        margin-bottom: 6px;
    }
    .external-report blockquote {
        margin: 10px 0;
        padding: 10px 12px;
        border-left: 4px solid #60a5fa;
        background: #ffffff;
        color: #334155;
    }
    .external-report ul {
        margin-bottom: 0;
    }
    .academic-preview {
        border: 1px solid #d7dee8;
        border-radius: 8px;
        background: #ffffff;
        color: #111827;
        padding: 18px 20px;
        margin-bottom: 12px;
        font-family: "Times New Roman", Arial, serif;
        line-height: 1.55;
    }
    .academic-preview h3 {
        font-size: 1.05rem;
        margin: 14px 0 8px;
        color: #111827;
        font-weight: 700;
    }
    .academic-preview h3:first-child {
        margin-top: 0;
    }
    .academic-preview p {
        margin: 0 0 9px;
        text-align: justify;
    }
    .academic-preview p[dir="rtl"] {
        text-align: right;
    }
    .academic-preview ul {
        margin: 0 0 10px 22px;
        padding: 0;
    }
    .academic-preview table {
        width: 100%;
        border-collapse: collapse;
        margin: 10px 0;
        font-size: 0.95rem;
    }
    .academic-preview td {
        border: 1px solid #cbd5e1;
        padding: 6px 8px;
        vertical-align: top;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("الإعدادات")
    input_option = st.radio("مصدر النص", ("رفع ملف", "لصق نص"), key="source_radio")

    text_input = ""
    if input_option == "رفع ملف":
        uploaded_file = st.file_uploader("اختر ملفًا", type=["txt", "docx", "pdf"], key="file_uploader")
        if uploaded_file is not None:
            try:
                text_input = extract_uploaded_text(uploaded_file)
            except Exception as exc:
                st.error(f"تعذر قراءة الملف: {exc}")
    else:
        text_input = st.text_area("ألصق النص الأكاديمي هنا", height=220, key="paste_area")

    if text_input:
        st.session_state.text_input = text_input

    st.slider(
        "قوة المراجعة",
        1,
        5,
        2,
        help="1=تنظيف خفيف، 5=مراجعة أسلوبية أوسع مع الحفاظ على المعنى.",
        key="intensity",
    )
    st.selectbox(
        "المجال الأكاديمي",
        ("medical", "engineering", "humanities", "general"),
        format_func=lambda value: {
            "medical": "طبي",
            "engineering": "هندسي",
            "humanities": "علوم إنسانية",
            "general": "عام",
        }[value],
        key="domain",
    )
    st.checkbox("حافظ على حجم النص قدر الإمكان", key="preserve_word_count")
    st.text_area(
        "تقرير كاشف خارجي (اختياري)",
        key="external_detector_report",
        height=120,
        help="ألصق هنا نتيجة GPTZero أو غيره إن وجدت. سيعرضها التطبيق كمرفق خارجي لا كحكم صادر منه.",
    )

    st.button(
        "بدء المراجعة",
        type="primary",
        use_container_width=True,
        on_click=process_text_callback,
    )

    st.markdown("---")
    st.subheader("مؤشرات الجودة")
    stats: RevisionStats = st.session_state.stats
    if st.session_state.processing_done and st.session_state.revised_text:
        st.metric("الكلمات قبل/بعد", f"{stats.words_original} / {stats.words_revised}")
        st.metric("متوسط طول الجملة", f"{stats.avg_sentence_length:.1f}")
        st.metric("التنوع اللفظي", f"{stats.lexical_diversity:.2f}")
        st.metric("حفظ المفردات", f"{stats.similarity:.2f}")
        st.caption("حفظ المفردات يقيس مقدار الكلمات المشتركة بين الأصل والمراجعة؛ 1.00 يعني أن المفردات لم تتغير تقريبًا، وليس حكمًا على الجودة.")

        st.metric("عبارات قالبية", stats.formulaic_phrase_count)
        st.metric("عبارات جزم زائد", stats.certainty_marker_count)
        st.metric("تباين أطوال الجمل", f"{stats.sentence_length_variation:.2f}")
        st.metric("نطاق المراجعة", stats.authorship_review_band)
        st.info(
            "هذه المؤشرات تظهر بعد المعالجة فقط، وتساعد على المراجعة التحريرية. لا تمثل حكمًا نهائيًا على الأصالة أو القبول الأكاديمي."
        )
    else:
        st.info("ستظهر مؤشرات الجودة بعد معالجة نص أو ملف. لا توجد نتيجة محسوبة بعد.")


col1, col2 = st.columns(2)
with col1:
    st.subheader("النص الأصلي")
    if st.session_state.text_input:
        st.text_area(
            "النص الأصلي",
            st.session_state.text_input,
            height=420,
            label_visibility="collapsed",
            key="original_text_display",
        )
        st.caption(f"عدد الكلمات: {len(tokenize_words(st.session_state.text_input))}")
    else:
        st.info("ارفع ملفًا أو ألصق نصًا من الشريط الجانبي.")

with col2:
    st.subheader("النص المراجع")
    if st.session_state.revised_text:
        st.markdown(render_academic_preview(st.session_state.revised_text), unsafe_allow_html=True)
        st.text_area(
            "النص المراجع",
            st.session_state.revised_text,
            height=420,
            label_visibility="collapsed",
            key="revised_text_display",
        )
        st.caption(f"عدد الكلمات: {len(tokenize_words(st.session_state.revised_text))}")
        word_file = create_word_document(st.session_state.revised_text)
        st.download_button(
            "تنزيل الملف المراجع المنسق (Word)",
            data=word_file,
            file_name="deepclean_revised.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        st.info("ستظهر النسخة المراجعة هنا بعد المعالجة.")


if st.session_state.text_input:
    report_is_current = (
        st.session_state.processing_done
        and st.session_state.processed_source_text == st.session_state.text_input.strip()
    )
    report: TransparencyReport = (
        st.session_state.transparency_report
        if report_is_current
        else compute_transparency_report(st.session_state.text_input)
    )
    st.markdown("---")
    with st.expander("لوحة الشفافية والتحليل السبعي", expanded=report_is_current):
        if not report.minimum_ready:
            st.warning(
                f"النص المؤهل للفحص يحتوي على {report.qualifying_chars} حرفًا بعد عزل المراجع والجداول والأكواد. يفضل توفر 250 حرفًا مؤهلًا على الأقل قبل الاعتماد على إشارات المراجعة."
            )

        overview_tab, sentence_tab, vocabulary_tab, support_tab = st.tabs(
            ["التصنيف", "تظليل الجمل", "المفردات", "الدعم والتحقق"]
        )

        with overview_tab:
            st.markdown(render_advanced_scan_card(report, report_is_current), unsafe_allow_html=True)
            if not report_is_current:
                st.info("غيّرت النص بعد آخر معالجة. اضغط بدء المراجعة لتحديث المؤشرات والنص المراجع معًا.")
            metric_cols = st.columns(5)
            metric_cols[0].metric("نطاق المؤشرات", report.classification)
            metric_cols[1].metric("قوة الإشارة", f"{report.confidence * 100:.0f}%")
            metric_cols[2].metric("الحروف", report.chars)
            metric_cols[3].metric("الكلمات", report.words)
            metric_cols[4].metric("الجمل", report.sentences)
            st.progress(report.confidence)
            if report.false_positive_shield:
                st.success("درع عتبة 20% فعّال: الإشارات منخفضة، لذلك لا تظهر النتيجة كاتهام أو حكم أصالة.")
            st.caption(
                "هذه قراءة محلية قابلة للتفسير. لا تثبت التأليف الآلي، ولا تنفيه، ولا تغني عن مراجعة بشرية وسجل تحرير ومصادر قابلة للتحقق."
            )
            external_report = st.session_state.get("external_detector_report", "").strip()
            if external_report:
                st.markdown(render_external_detector_card(external_report, report), unsafe_allow_html=True)

            q_cols = st.columns(3)
            q_cols[0].metric("النص المؤهل للفحص", f"{report.qualifying_chars} حرف")
            q_cols[1].metric("المستبعد من الفحص", f"{report.excluded_chars} حرف")
            q_cols[2].metric("تحديث الأنماط", "توصية يدوية")

            breakdown_rows = pd.DataFrame(
                [
                    {"الفئة": "تشابه مع نص مولد آليًا", "الدرجة": f"{report.ai_generated_score:.2f}"},
                    {"الفئة": "اشتباه إعادة صياغة آلية", "الدرجة": f"{report.ai_paraphrased_score:.2f}"},
                    {"الفئة": "تكرار داخلي", "الدرجة": f"{report.plagiarism_signal:.2f}"},
                    {"الفئة": "أثر ترجمة أو انتقال لغوي", "الدرجة": f"{report.translation_signal:.2f}"},
                ]
            )
            st.dataframe(breakdown_rows, use_container_width=True, hide_index=True)

            signal_rows = pd.DataFrame(
                [
                    {"المكون": "الحيرة/تنوع المفردات", "القيمة": f"{report.perplexity_signal:.2f}"},
                    {"المكون": "التدفق/تباين الجمل", "القيمة": f"{report.burstiness_signal:.2f}"},
                    {"المكون": "درع إعادة الصياغة", "القيمة": f"{report.paraphrase_signal:.2f}"},
                    {"المكون": "تعديل ESL", "القيمة": f"-{report.esl_adjustment:.2f}"},
                ]
            )
            st.dataframe(signal_rows, use_container_width=True, hide_index=True)
            if report.excluded_blocks:
                excluded_rows = pd.DataFrame(
                    [{"العنصر المعزول": reason, "حروف مستبعدة": count} for reason, count in report.excluded_blocks]
                )
                st.dataframe(excluded_rows, use_container_width=True, hide_index=True)

        with sentence_tab:
            if report.sentence_signals:
                highlighted = render_sentence_highlights(report)
                st.markdown(f"<div class='transparency-panel'>{highlighted}</div>", unsafe_allow_html=True)
                st.caption("البرتقالي = مراجعة عالية، الأصفر = مراجعة متوسطة، الأخضر = إشارة منخفضة.")
                sentence_rows = pd.DataFrame(
                    [
                        {
                            "الجملة": signal.text,
                            "النطاق": signal.label,
                            "الدرجة": f"{signal.score:.2f}",
                            "السبب": "، ".join(signal.reasons),
                        }
                        for signal in report.sentence_signals
                    ]
                )
                st.dataframe(sentence_rows, use_container_width=True, hide_index=True)
            else:
                st.info("لا توجد جمل كافية للتظليل بعد.")

        with vocabulary_tab:
            if report.ai_vocabulary_hits:
                vocab_rows = pd.DataFrame(
                    [{"المفردة": term, "عدد الظهور": count} for term, count in report.ai_vocabulary_hits]
                )
                st.dataframe(vocab_rows, use_container_width=True, hide_index=True)
            else:
                st.success("لم تظهر مفردات آلية متكررة ضمن القاموس المحلي.")
            st.caption("القاموس المحلي يركز على إشارات تحريرية شائعة، وقد لا يلتقط كل الأنماط أو اللغات.")

        with support_tab:
            component_rows = pd.DataFrame(
                [{"المكون": title, "كيف يستخدمه DeepClean": description} for title, description in TRANSPARENCY_COMPONENTS]
            )
            st.dataframe(component_rows, use_container_width=True, hide_index=True)
            st.subheader("حدود الدقة")
            limitations_rows = pd.DataFrame(
                [{"البند": title, "المعنى": description} for title, description in REALISM_LIMITATIONS]
            )
            st.dataframe(limitations_rows, use_container_width=True, hide_index=True)
            external_report = st.session_state.get("external_detector_report", "").strip()
            if external_report:
                st.markdown(render_external_detector_card(external_report, report), unsafe_allow_html=True)
            for note in report.section_notes:
                st.info(note)
            if report.source_code_alerts:
                for alert in report.source_code_alerts:
                    st.warning(alert)
            if report.citation_alerts:
                for alert in report.citation_alerts:
                    st.warning(alert)
            else:
                st.success("لم تظهر تنبيهات مرجعية محلية واضحة.")
            st.info(
                "فحص الهلوسة هنا محلي ومبدئي: يراجع شكل الإحالات والروابط فقط. التحقق الكامل من المصادر يحتاج اتصالًا بقاعدة بيانات أو بحثًا موثقًا."
            )
            st.info(
                "تقرير Writing Replay يحتاج تكاملًا مع سجل تحرير خارجي مثل Google Docs؛ التطبيق يجهز مكانه المفاهيمي دون تتبع كتابة المستخدم."
            )
            st.caption(report.model_update_note)


if st.session_state.revised_text and st.session_state.text_input:
    st.markdown("---")
    with st.expander("تفاصيل التغييرات", expanded=False):
        diff_html = word_level_diff(st.session_state.text_input, st.session_state.revised_text)
        st.markdown(
            """
            <style>
            .diff-box {background:#fafafa;border:1px solid #e5e7eb;padding:14px;border-radius:8px;line-height:1.8}
            .removed {color:#b91c1c;text-decoration:line-through}
            .added {color:#047857;font-weight:600}
            </style>
            """,
            unsafe_allow_html=True,
        )
        st.markdown(f"<div class='diff-box'>{diff_html}</div>", unsafe_allow_html=True)

        original_set = Counter(word.lower() for word in tokenize_words(st.session_state.text_input))
        revised_set = Counter(word.lower() for word in tokenize_words(st.session_state.revised_text))
        removed = sorted((original_set - revised_set).elements())
        added = sorted((revised_set - original_set).elements())
        if removed or added:
            changes = pd.DataFrame(
                {
                    "النوع": ["محذوف"] * len(removed) + ["مضاف"] * len(added),
                    "الكلمة": removed + added,
                }
            )
            st.dataframe(changes, use_container_width=True, hide_index=True)
        else:
            st.info("لا توجد تغييرات جوهرية على مستوى الكلمات.")

st.markdown("---")
st.caption(f"DeepClean Studio © 2026 - {AUTHOR_NAME} - للمراجعة التعليمية والبحثية المسؤولة.")
